from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/fanzengxia/Documents/GitHub/frontmind-dashboard")
OUTPUT = ROOT / "outputs/kp_youth_2026/FrontMind_深圳鲲鹏青年项目申请简介_纯文本段落版.docx"

FONT = "Arial Unicode MS"
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(90, 90, 90)


def set_font(run, size=None, bold=None, color=BLACK):
    run.font.name = FONT
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = color


def set_style_font(style, size, bold=False, color=BLACK):
    style.font.name = FONT
    r_fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = color


def add_page_field(paragraph):
    run = paragraph.add_run()
    set_font(run, size=9, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def add_body(doc, text, *, indent=True):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.2
    if indent:
        paragraph.paragraph_format.first_line_indent = Pt(22)
    run = paragraph.add_run(text)
    set_font(run, size=11)
    return paragraph


def add_section(doc, heading, paragraphs):
    paragraph = doc.add_paragraph(style="Heading 1")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(heading)
    set_font(run, size=15, bold=True)
    for text in paragraphs:
        add_body(doc, text)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, 11, False)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.2

    heading = doc.styles["Heading 1"]
    set_style_font(heading, 15, True)
    heading.paragraph_format.space_before = Pt(16)
    heading.paragraph_format.space_after = Pt(6)
    heading.paragraph_format.keep_with_next = True

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.paragraph_format.space_after = Pt(0)
    add_page_field(footer_paragraph)

    properties = doc.core_properties
    properties.title = "FrontMind 深圳鲲鹏青年创新创业项目申请简介（纯文本段落版）"
    properties.subject = "深圳鲲鹏青年创新创业项目申报材料"
    properties.author = "FrontMind"
    properties.keywords = "FrontMind, 深圳鲲鹏青年, 人工智能, GEO, 可信智能体"
    properties.comments = "纯文本段落版；经营与权属数据中的待补内容须在提交前核验。"


def add_title_block(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("深圳鲲鹏青年创新创业项目申请简介")
    set_font(run, size=20, bold=True)

    project = doc.add_paragraph()
    project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    project.paragraph_format.space_after = Pt(14)
    run = project.add_run("FrontMind：可信智能体驱动的企业 GEO 与 AI 增长平台")
    set_font(run, size=14, bold=True)

    metadata = [
        "申报主体：深圳市超前无限科技有限公司（最终以营业执照为准）",
        "项目负责人：夏凡增",
        "所属领域：人工智能（主赛道），软件和信息服务（关联支撑）",
    ]
    for text in metadata:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(text)
        set_font(run, size=10.5, color=GRAY)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)

    add_body(
        doc,
        "FrontMind 源自香港中文大学（深圳）智能决策实验室的研究与成果转化实践，面向生成式人工智能正在重塑企业“发现、比较、决策、推荐”链路的新趋势，建设可信智能体驱动的企业 GEO（Generative Engine Optimization，生成式引擎优化）与 AI 增长平台。项目以企业知识和可信证据为底座，通过跨平台监测、品牌语义诊断、问答与内容资产建设、增长触达和企业级工作流部署，帮助企业在生成式人工智能环境中形成准确、稳定、可验证、可持续优化的品牌与产品认知。",
    )


def main():
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    add_section(
        doc,
        "一、产业导向与商业模式",
        [
            "本项目主匹配深圳市“20+8”产业集群中的人工智能产业集群，并与软件和信息服务产业集群形成交叉支撑。FrontMind 的核心能力包括生成式人工智能评测、企业知识治理、可信智能体、跨模型监测、智能决策和企业级软件平台，属于人工智能技术在现代服务业和企业数字化领域的产业化应用。同时，项目可服务深圳先进制造、科技服务、消费品牌和出海企业，将企业的技术、产品、资质和场景优势转化为可被主流生成式人工智能准确理解和引用的数字资产，具有明显的产业赋能属性。",
            "生成式人工智能正在成为用户获取信息、比较产品和形成购买决策的重要入口，但企业普遍缺少面向该入口的系统化经营能力。首先，企业官网、产品资料、媒体报道、案例和内部知识彼此分散，大模型容易出现事实缺失、信息过时、品牌定位偏差和引用来源不足。其次，传统 SEO 主要关注网页排名，无法充分衡量企业在大模型回答中的提及率、排序、推荐倾向、情绪、引用来源和竞争品牌占位。再次，单次提问结果具有随机性，企业缺少跨平台、重复采样和可追溯证据，难以判断优化是否真实有效。最后，市场上的监测、内容生产、渠道分发、销售线索和企业系统往往相互割裂，诊断结果难以转化为持续执行和复测闭环。",
            "针对上述问题，FrontMind 形成三个相互衔接的产品和服务层。MindPromise 面向企业可信品牌认知，提供品牌事实图谱、证据库、品牌语义资产审计、问题组合设计、多平台重复监测、引用与情绪分析、竞品占位分析、问答树和内容资产建设。现有监测框架覆盖豆包、腾讯元宝、DeepSeek、百度 AI+、通义千问、Kimi 等主流中文平台，并通过多次采样降低单次回答波动。MindReach 面向市场意图和增长转化，识别用户需求信号，设计回答逻辑、内容和渠道任务，连接线索触达与效果回流。MindNexus/FDE 面向企业级部署，将私有知识、业务流程、API、权限、审计和安全机制接入企业生产环境。",
            "项目采用“标准化软件订阅、专业服务交付、企业级定制部署”相结合的商业模式。标准化产品以单问题诊断、多平台监测、知识库、问答体系和季度优化服务作为低门槛入口，按月或按季度形成持续订阅收入；专业服务围绕品牌定位、证据治理、内容资产和渠道执行形成项目收入；对于知识复杂、系统众多或数据安全要求较高的客户，则通过 FDE、私有化部署、系统集成和联合研发提高客单价。项目还将与高校、园区、行业协会、企业服务机构和产业链龙头建立联合诊断、联合交付与客户转介机制，降低获客成本并扩大行业覆盖。",
            "项目的落地路径是先通过可量化诊断确认企业在生成式人工智能中的认知问题，再以知识库和持续监测建立订阅关系，随后延伸到内容、渠道、线索和企业系统部署。随着平台能力、行业知识包和自动化交付流程逐步成熟，重复性工作将由系统承担，专家团队聚焦高价值策略和复杂企业部署，从而提升交付效率、毛利率和客户续费率，形成可规模化、可持续的商业闭环。",
        ],
    )

    add_section(
        doc,
        "二、核心团队与创业实践",
        [
            "项目负责人夏凡增为香港中文大学（深圳）数据科学方向博士研究生，研究聚焦 GEO 与营销智能体、上下文决策和稳健大模型应用，拥有纽约大学计算机科学硕士背景，本科阶段接受计算机工程训练。公开履历显示，其曾在高校参与自然语言处理算法研究，并具有亚马逊云科技全栈工程实践，具备从算法研究、云端工程到企业产品化的复合能力。自 2026 年起，夏凡增担任 FrontMind 联合创始人兼 CEO，主要负责项目战略、产品路线、核心技术、融资、产业合作和关键客户拓展。",
            "项目依托香港中文大学（深圳）智能决策实验室在可信人工智能、在线决策、强化学习、控制和 GEO 等方向的研究基础。实验室公开资料显示，其成立于 2022 年，拥有 14 名研究成员，累计发表 52 篇论文并拥有 4 项专利，成员具有清华大学、加州理工学院、纽约大学、柏林工业大学和上海交通大学等教育背景，以及亚马逊、谷歌、字节跳动、华为等产业经历。上述资源为项目提供研究方法、人才和产业问题理解基础，但实验室论文、专利和合作不直接等同于申请企业资产；如在正式申报中列入企业知识产权或企业合作成果，需提供许可、转让、共同研发或合同主体证明。",
            "项目团队按照算法研究、平台工程、品牌策略与内容、商务交付与客户成功四类能力进行分工。算法研究人员负责 GEO 评测、可信智能体、跨模型实验和决策优化；平台工程人员负责前后端系统、数据架构、账号权限、支付、部署和运维；品牌策略与内容人员负责企业事实治理、问答架构、内容标准和行业交付；商务交付与客户成功人员负责客户拓展、伙伴合作、项目管理、续费和服务质量。除负责人外，正式申报所需核心成员姓名、学历、专业背景、行业经历、劳动或合作关系及具体职责为【待补】，并应确保团队人数及材料符合管理办法要求。",
            "在科技成果转化方面，项目已将研究问题转化为可在线访问的品牌网站、企业服务平台和产品模块，建立了从企业知识、问题选择、意图优化、回答逻辑、多平台监测到内容资产和渠道交付的业务流程。项目代码已经具备多租户权限、企业知识库、监测看板、合同、支付、账号开通、服务进度和操作审计等生产化能力，说明团队不仅具备算法研究能力，也具备将研究原型转化为企业软件和交付流程的工程能力。",
            "公开资料显示，相关团队曾获得 2024 年“未来香港”创新科技大赛冠军，并在 2025 年中国创新创业大赛、上海海聚英才全球创新创业大赛、C-Star 等创新创业活动中获奖或入选；实验室还公开展示了在电力、信息通信等领域的研究合作实践。这些经历体现了团队识别产业问题、组织跨学科研发和进行成果展示的能力。正式申报时，应根据获奖证书、合作协议和合同主体，分别说明个人、实验室、项目团队与申请企业之间的关系，避免将实验室合作直接写成企业订单。",
        ],
    )

    add_section(
        doc,
        "三、创新能力与市场潜力",
        [
            "FrontMind 的核心竞争优势不是单一提示词或内容生成能力，而是由可信证据、跨平台评测、智能决策、执行复测和企业工程组成的连续系统。项目首先将企业的产品主张、技术参数、资质、案例和媒体信息映射到具体来源、时效和权威等级，并区分得到支持、存在冲突、被遗漏和无法核验等状态；随后通过跨模型、跨问题、多次采样测量品牌提及率、排序、情绪、引用来源和竞品占位，降低单次回答噪声；再依据结果确定问题、内容和渠道动作的优先级，并在执行后复测，形成可验证的持续优化闭环。",
            "项目的技术壁垒主要体现在四个方面。第一，证据型企业知识体系能够把事实、主张和权威来源连接起来，提高大模型回答的可验证性。第二，跨平台重复采样和统一指标体系可以形成行业基线，支持不同模型、不同问题和不同时间的纵向与横向比较。第三，项目结合上下文决策和稳健智能体方法，对问题组合、内容资产和渠道动作进行优先级排序，使资源投入与认知差距相匹配。第四，多租户权限、密钥保护、操作审计、支付、账号开通和系统集成等工程能力，使项目可以从诊断工具进一步进入企业生产环境。",
            "项目已形成自主研发的平台源代码、数据结构、评测流程、行业问答方法、交付模板和智能体工作流。申请企业现有软件著作权、专利、商标、域名权属以及核心成员职务成果协议情况为【待补】。正式申报时应提供知识产权清单、代码版本记录、域名和产品页面权属、核心成员保密与职务成果协议、第三方模型和开源组件许可清单；对于实验室成果，则应提供专利权人证明以及许可、转让或共同研发文件。在缺少权属文件的情况下，不应将实验室 4 项专利直接表述为企业自有知识产权。",
            "与传统 SEO 和内容服务相比，FrontMind 直接以生成式人工智能回答、推荐和引用为优化对象，能够处理传统网页排名无法覆盖的模型认知问题。与 Profound、Scrunch、Peec 等 AI 可见度监测产品相比，FrontMind 更强调中文主流人工智能平台、企业可信证据、品牌知识库、问答与内容资产、渠道执行以及复测闭环。与通用 AI 咨询或定制开发相比，FrontMind 通过标准化诊断和订阅平台提高复用能力，同时保留 FDE 和私有部署能力以满足复杂客户需求。项目的差异化价值在于将“看见问题、解释原因、执行优化、验证结果”连接为一体。",
            "当前项目已经上线 FrontMind 官方网站、产品说明和在线服务入口，形成 MindPromise、MindReach、MindNexus/FDE 的产品架构；已经实现六类中文主流人工智能平台的监测框架，以及知识库、问题选择、意图优化、回答逻辑、持续监测、渠道分发、进度报告和内容资产等模块；并建立了从服务选择、合同与支付、账号开通到项目交付的业务闭环基础。FrontMind 官网公开披露服务“100+ 行业客户/品牌案例”，正式申报时须以合同、订单、发票、付款记录或经授权案例清单核验，不能仅以公开展示作为已实现营收依据。",
            "市场需求正在快速形成。截至 2025 年末，我国生成式人工智能用户规模达到 6.02 亿，普及率为 42.8%，完成备案的生成式人工智能服务超过 750 款；2025 年我国人工智能核心产业规模超过 1.2 万亿元。深圳 2025 年人工智能核心产业规模超过 2200 亿元，规模以上人工智能企业超过 2600 家，并拥有约 2.6 万家国家高新技术企业。上述企业均面临产品知识、品牌认知、客户问答和市场渠道向生成式人工智能迁移的需求。",
            "由于 GEO 尚缺少统一的官方市场统计口径，项目采用自下而上的方式测算初始市场。以深圳约 2.6 万家国家高新技术企业为潜在客户池，若首阶段转化其中 1%，按年度综合客单价 3 万元计算，对应约 780 万元的年服务规模。该数字仅为经营假设，不代表已获得订单或收入。未来项目可从深圳拓展到粤港澳大湾区、全国科技企业和出海品牌，并通过行业知识包、伙伴渠道、订阅续费和企业级部署持续扩大客户数量和综合客单价。",
        ],
    )

    add_section(
        doc,
        "四、运营及造血能力",
        [
            "申请主体拟为深圳市超前无限科技有限公司，企业法定名称、统一社会信用代码、成立日期、注册地址和股权结构须以营业执照最终确认。项目现已具备独立品牌、公开网站、产品架构和企业服务平台，并形成从诊断、合同、支付、账号开通到持续交付的流程基础。知识库、监测、权限、审计、服务订单和交付进度已经纳入统一系统，为后续标准化销售、客户续费和规模化交付提供了运营基础。",
            "企业财务和经营数据尚未随本次材料完整提供。正式申报时应补充：截至 2026 年【待补月份】企业累计营业收入【待补】万元、2026 年当年营业收入【待补】万元、已回款【待补】万元、经营成本【待补】万元和现金余额【待补】万元；现有在手订单【待补】份，合同含税金额【待补】万元，其中已交付【待补】万元、待交付【待补】万元；主要客户、合同周期和回款状态应与合同、发票、银行流水和验收材料保持一致。公开案例、意向客户、实验室合作和企业已签订单必须分别统计，只有合同主体为申请企业且金额和履约状态可以验证的项目，方可列入在手订单或营业收入。",
            "融资方面，企业当前股东投入和团队自筹资金合计为【待补】万元，可支持现阶段【待补】个月的研发和运营；外部融资进度为【待补：尚未融资、正在接洽或已完成具体轮次及金额】。如已获得投资，应提供投资协议、付款凭证和工商变更信息；如尚未完成外部融资，则应说明创始团队出资、客户回款和经营现金流对项目研发与交付的支持能力。企业现有合作协议资源为【待补】项，涉及高校、园区、行业机构、客户或产业伙伴的合作，应逐项核验协议主体、期限、金额、具体内容和履约状态。",
            "项目的造血逻辑来自三个层次。第一，以低门槛诊断和标准化监测获取客户，通过月度或季度订阅形成稳定续费收入。第二，以知识库、问答、内容资产和渠道执行扩大单客户服务范围，提高复购和综合客单价。第三，以企业 FDE、私有化部署、系统集成和联合研发服务复杂客户，形成更高价值的项目收入。随着行业知识包和交付自动化程度提高，单位客户交付成本将逐步下降，订阅收入占比和毛利率有望提升，从而减少对一次性项目和外部融资的依赖。",
            "结合当前产品基础和市场拓展计划，项目建议将 2027 年、2028 年和 2029 年新增产值目标分别设定为 300 万元、800 万元和 1800 万元，三年累计新增产值 2900 万元。对应的付费客户目标分别为 100 家、200 家和 360 家，综合年客单价假设分别为 3 万元、4 万元和 5 万元。上述数据属于规划目标，须由企业负责人结合实际销售周期、交付能力、成本结构和现金流进一步确认，不得表述为已实现收入。",
            "项目预计在三年内累计新增就业 35 人，其中 2027 年新增约 8 人，2028 年累计新增约 18 人，2029 年累计新增约 35 人。新增岗位主要包括算法研究、平台工程、数据与知识运营、品牌策略与内容、行业交付、商务拓展和客户成功。项目将优先吸纳具有人工智能、计算机、数据科学、营销科技和行业数字化背景的青年人才，并通过高校科研合作和产业实践形成稳定的人才培养与转化机制。",
            "项目资金将优先用于证据型知识引擎、跨平台评测、可信智能体决策、产品安全、知识产权确权、行业知识包和交付自动化建设。短期目标是完成核心产品版本、测试数据集和首批可核验标杆客户；中期目标是形成稳定续费、缩短交付周期并提高订阅收入占比；长期目标是发展企业级部署和全国及出海市场，形成多行业复制能力和正向经营现金流。根据管理办法，推荐或邀请入库项目可获得最高 5 万元一次性创新奖励，后续符合条件且在深圳实际运营的项目可另行申请不超过研发投入 50%、最高 100 万元的研发补助。相关政策资金不计入企业已实现营收，应按对应条件和程序单独申请。",
        ],
    )

    note = doc.add_paragraph(style="Normal")
    note.paragraph_format.space_before = Pt(12)
    note.paragraph_format.space_after = Pt(0)
    note.paragraph_format.line_spacing = 1.2
    run = note.add_run(
        "说明：文中所有【待补】内容及公开披露数据，须在正式提交前以营业执照、身份证明、学历证明、团队材料、财务报表、纳税记录、银行流水、合同、发票、回款、合作协议、获奖证书和知识产权权属文件进行核验。"
    )
    set_font(run, size=10, color=GRAY)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
