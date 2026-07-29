from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/fanzengxia/Documents/GitHub/frontmind-dashboard")
OUTPUT = ROOT / "outputs/kp_youth_2026/FrontMind_深圳鲲鹏青年项目申请简介_两页精简版.docx"

FONT = "Arial Unicode MS"
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(88, 88, 88)


def set_font(run, size=None, bold=None, color=BLACK):
    run.font.name = FONT
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = color


def set_style_font(style, size, bold=False, color=BLACK):
    style.font.name = FONT
    r_fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = color


def add_page_number(paragraph):
    run = paragraph.add_run()
    set_font(run, size=9, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def add_body(doc, text, *, size=10.2, gray=False, indent=True):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2.5)
    paragraph.paragraph_format.line_spacing = 1.05
    if indent:
        paragraph.paragraph_format.first_line_indent = Pt(21)
    run = paragraph.add_run(text)
    set_font(run, size=size, color=GRAY if gray else BLACK)
    return paragraph


def add_section(doc, heading, paragraphs):
    paragraph = doc.add_paragraph(style="Heading 1")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(heading)
    set_font(run, size=13.2, bold=True)
    for text in paragraphs:
        add_body(doc, text)


def add_page_break(doc):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.38)

    normal = doc.styles["Normal"]
    set_style_font(normal, 10.2)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.line_spacing = 1.05

    heading = doc.styles["Heading 1"]
    set_style_font(heading, 13.2, bold=True)
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(3)
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.keep_together = True

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.paragraph_format.space_after = Pt(0)
    add_page_number(footer_paragraph)

    properties = doc.core_properties
    properties.title = "FrontMind 深圳鲲鹏青年创新创业项目申请简介（两页精简版）"
    properties.subject = "深圳鲲鹏青年创新创业项目申报材料"
    properties.author = "FrontMind"
    properties.keywords = "FrontMind, 深圳鲲鹏青年, 人工智能, GEO, 可信智能体"
    properties.comments = "两页精简版；待补及公开披露数据须在正式提交前核验。"


def add_title(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    set_font(title.add_run("深圳鲲鹏青年创新创业项目申请简介"), size=18, bold=True)

    project = doc.add_paragraph()
    project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    project.paragraph_format.space_after = Pt(4)
    set_font(
        project.add_run("FrontMind：可信智能体驱动的企业 GEO 与 AI 增长平台"),
        size=13,
        bold=True,
    )

    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_after = Pt(5)
    set_font(
        metadata.add_run(
            "申报主体：深圳市超前无限科技有限公司（以营业执照为准）｜"
            "项目负责人：夏凡增｜所属领域：人工智能"
        ),
        size=9.2,
        color=GRAY,
    )


def main():
    doc = Document()
    configure_document(doc)
    add_title(doc)

    add_section(
        doc,
        "一、产业导向与商业模式",
        [
            "FrontMind 主匹配深圳“20+8”产业集群中的人工智能产业集群，并与软件和信息服务业交叉融合。生成式人工智能正成为用户发现、比较和选择产品的重要入口，但企业资料分散、事实过时、引用不足，传统 SEO 又无法衡量品牌在大模型回答中的提及率、推荐倾向、情绪、信源和竞品占位；单次回答具有随机性，监测、内容、渠道与线索系统割裂，企业迫切需要可验证、可持续的智能认知经营工具。",
            "项目以可信企业知识和证据为底座，形成 MindPromise、MindReach、MindNexus/FDE 三层产品。MindPromise 提供品牌事实图谱、证据库、语义资产审计及豆包、腾讯元宝、DeepSeek、百度 AI+、通义千问、Kimi 等主流平台的重复监测，输出引用、情绪、竞品与问答内容优化结果；MindReach 识别用户意图，连接内容、渠道和线索触达；MindNexus/FDE 将私有知识、业务流程、API、权限、审计和安全机制接入企业生产环境。",
            "商业模式为“标准化软件订阅＋专业服务＋企业级部署”。项目以低门槛诊断切入，以月度或季度监测、知识库和持续优化形成订阅收入，以品牌定位、证据治理、内容资产和渠道执行形成专业服务收入，并通过 FDE、私有化部署、系统集成和联合研发提升复杂客户客单价。项目将与高校、园区、行业协会和企业服务机构开展联合诊断、交付和客户转介，形成“诊断—订阅—深度部署—持续复测”的规模化闭环。",
        ],
    )

    add_section(
        doc,
        "二、核心团队与创业实践",
        [
            "项目负责人夏凡增为香港中文大学（深圳）数据科学方向博士研究生，研究聚焦 GEO 与营销智能体、上下文决策及稳健大模型应用，拥有纽约大学计算机科学硕士背景，并具有自然语言处理算法研究、亚马逊云科技全栈工程和企业产品化经验。其现任 FrontMind 联合创始人兼 CEO，负责战略、产品路线、核心技术、融资、产业合作和关键客户拓展。",
            "项目依托香港中文大学（深圳）智能决策实验室在可信人工智能、在线决策、强化学习、控制和 GEO 等方向的研究基础。实验室公开资料显示，团队成立于 2022 年，拥有 14 名成员、52 篇论文和 4 项专利。项目团队按算法研究、平台工程、品牌策略与内容、商务交付与客户成功分工，已将研究方法转化为可访问的网站和企业服务平台，并实现知识库、监测看板、多租户权限、合同支付、账号开通、服务进度及操作审计等生产化能力，体现了从科研原型到软件产品和交付流程的转化能力。",
            "公开资料显示，相关团队曾获 2024 年“未来香港”创新科技大赛冠军，并在 2025 年中国创新创业大赛、上海海聚英才全球创新创业大赛、C-Star 等活动中获奖或入选。正式申报时须补充其他核心成员姓名、学历、经历、职责及劳动或合作关系，并以证书、协议核验奖项和成果主体；实验室论文、专利和合作仅在取得许可、转让或共同研发证明后，方可作为申请企业成果。",
        ],
    )

    add_page_break(doc)

    add_section(
        doc,
        "三、创新能力与市场潜力",
        [
            "FrontMind 的壁垒不是单一提示词或内容生成，而是“可信证据—跨平台评测—智能决策—执行复测—企业工程”的连续系统。项目把产品主张、技术参数、资质和案例映射到具体来源、时效与权威等级；通过跨模型、跨问题、多次采样统一测量提及、排序、情绪、引用和竞品占位；再依据认知差距安排问题、内容和渠道动作，并在执行后复测，形成可解释、可验证的增长闭环。多租户权限、密钥保护、操作审计和系统集成能力，使产品能够进入企业生产环境。",
            "相较传统 SEO 和内容服务，FrontMind 直接优化生成式人工智能回答、推荐与引用；相较 Profound、Scrunch、Peec 等可见度监测产品，更强调中文主流平台、可信证据、企业知识库以及从诊断到执行的闭环；相较通用 AI 咨询，又以标准化平台提高复用效率，并保留 FDE 和私有部署能力。项目已形成自主研发的平台代码、数据结构、评测流程、行业问答方法和智能体工作流，企业软件著作权、专利、商标及职务成果协议为【待补】，实验室专利未经权属证明不得列作企业自有知识产权。",
            "目前项目已上线 FrontMind 官网和在线服务入口，形成三层产品架构及六类中文主流人工智能平台监测框架，具备知识库、问题选择、意图优化、回答逻辑、持续监测、渠道分发、内容资产、合同支付与项目交付功能。官网公开披露服务“100+行业客户/品牌案例”，正式申报时应以企业名下合同、发票、回款或经授权案例清单核验，不将公开展示等同于已实现收入。",
            "截至 2025 年末，我国生成式人工智能用户约 6.02 亿、普及率 42.8%，人工智能核心产业规模超过 1.2 万亿元；深圳人工智能核心产业规模超过 2200 亿元，规模以上人工智能企业超过 2600 家，并拥有约 2.6 万家国家高新技术企业。GEO 尚无统一官方市场口径，按深圳 2.6 万家高新技术企业中首阶段转化 1%、年度综合客单价 3 万元测算，初始年服务机会约 780 万元。该数值为经营假设，未来可向大湾区、全国科技企业和出海品牌扩展。",
        ],
    )

    add_section(
        doc,
        "四、运营及造血能力",
        [
            "项目已具备独立品牌、公开网站、产品平台和从诊断、合同、支付、账号开通到持续交付的运营基础。企业真实经营数据须据实补充：截至 2026 年【月份】，累计营收【】万元、已回款【】万元、在手订单【】份／【】万元、现金余额【】万元；股东及团队自筹【】万元，外部融资【尚未融资／轮次及金额】，企业名下合作协议【】项。上述内容须与财务报表、纳税记录、银行流水、合同、发票、回款和验收材料一致，意向客户、公开案例及实验室合作不计作企业订单。",
            "项目的造血路径分为三层：标准化诊断和持续监测形成订阅续费，知识库、问答、内容和渠道服务提高复购与综合客单价，FDE、私有化部署和系统集成形成高价值项目收入。随着行业知识包沉淀和交付自动化，单位客户成本将下降，订阅收入与毛利率有望提高，从而逐步减少对一次性项目和外部融资的依赖。",
            "结合当前产品基础，项目建议将 2027—2029 年新增产值目标设为 300 万元、800 万元和 1800 万元，三年累计 2900 万元；对应付费客户目标为 100 家、200 家和 360 家。三年计划累计新增就业约 35 人，其中 2027 年新增 8 人、2028 年累计 18 人、2029 年累计 35 人，岗位覆盖算法、平台工程、数据与知识运营、品牌策略、行业交付、商务拓展和客户成功。上述均为规划目标，须由企业依据实际销售周期、成本结构和现金流确认。",
            "项目资金将重点投入证据型知识引擎、跨平台评测、可信智能体决策、产品安全、知识产权确权、行业知识包和交付自动化，短期完成核心版本与首批可核验标杆客户，中期提升续费和订阅收入占比，长期拓展企业级部署、全国及出海市场。文中所有【待补】及公开披露数据，提交前须以营业执照、身份与学历证明、团队材料、财务与税务资料、合同回款、合作协议、获奖证书和知识产权权属文件逐项核验。",
        ],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
