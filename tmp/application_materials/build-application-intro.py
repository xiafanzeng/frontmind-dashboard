from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/fanzengxia/Documents/GitHub/frontmind-dashboard")
OUT = ROOT / "outputs/kp_youth_2026/FrontMind_深圳鲲鹏青年项目申请简介_草拟版.docx"

NAVY = "19253F"
PURPLE = "5C4BD9"
PURPLE_LIGHT = "EEEAFE"
GOLD = "C6A452"
TEXT = "222A35"
MUTED = "647084"
LINE = "D8DEE8"
PALE_BLUE = "EAF2FB"
PALE_YELLOW = "FFF2CC"
PALE_GREEN = "E8F3EC"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size=6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_width(table, width_twips=9360, indent_twips=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_twips))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph, enabled=True) -> None:
    paragraph.paragraph_format.keep_with_next = enabled


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_field(run, instruction: str) -> None:
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_hyperlink(paragraph, text: str, url: str, color=PURPLE) -> None:
    part = paragraph.part
    rid = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_color = OxmlElement("w:color")
    r_color.set(qn("w:val"), color)
    r_pr.append(r_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_font_family(r_fonts, family="Arial Unicode MS") -> None:
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), family)


def set_run_fonts(run, family="Arial Unicode MS") -> None:
    run.font.name = family
    set_font_family(run._element.rPr.rFonts, family)


def add_text(
    paragraph,
    text: str,
    *,
    bold=False,
    italic=False,
    color=TEXT,
    size=None,
    highlight=None,
):
    run = paragraph.add_run(text)
    set_run_fonts(run)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    if highlight:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), highlight)
        run._r.get_or_add_rPr().append(shd)
    return run


def add_body(doc, text: str, *, bold_lead: str | None = None, style=None, after=6):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        add_text(p, bold_lead, bold=True)
        add_text(p, text[len(bold_lead) :])
    else:
        add_text(p, text)
    return p


def add_bullet(doc, text: str, *, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.18
    add_text(p, text)
    return p


def add_numbered(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.18
    add_text(p, text)
    return p


def add_heading(doc, text: str, level=1):
    p = doc.add_heading(text, level=level)
    keep_with_next(p)
    return p


def add_callout(doc, title: str, body: str, fill=PALE_YELLOW, title_color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    set_table_borders(table, color=GOLD, size=10)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, start=180, bottom=130, end=180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_text(p, title, bold=True, color=title_color, size=11)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.18
    add_text(p2, body)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_table(
    doc,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
    *,
    header_fill=NAVY,
    first_col_fill=None,
    font_size=9,
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, sum(widths), 120)
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    keep_row_together(hdr)
    for index, (cell, header, width) in enumerate(zip(hdr.cells, headers, widths)):
        set_cell_width(cell, width)
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_text(p, header, bold=True, color=WHITE, size=font_size)
    for row_index, row_values in enumerate(rows):
        row = table.add_row()
        keep_row_together(row)
        for col_index, (cell, value, width) in enumerate(zip(row.cells, row_values, widths)):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            if first_col_fill and col_index == 0:
                set_cell_shading(cell, first_col_fill)
            elif row_index % 2 == 1:
                set_cell_shading(cell, "F7F9FC")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_text(p, value, bold=(first_col_fill and col_index == 0), size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    set_font_family(normal._element.rPr.rFonts)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[style_name]
        style.font.name = "Arial Unicode MS"
        set_font_family(style._element.rPr.rFonts)
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor.from_string(TEXT)

    for name, size, color, before, after in (
        ("Title", 25, NAVY, 0, 12),
        ("Subtitle", 13, MUTED, 0, 10),
        ("Heading 1", 16, NAVY, 18, 8),
        ("Heading 2", 13, PURPLE, 12, 5),
        ("Heading 3", 11, NAVY, 8, 3),
    ):
        style = styles[name]
        style.font.name = "Arial Unicode MS"
        set_font_family(style._element.rPr.rFonts)
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Small Note" not in styles:
        note = styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        note = styles["Small Note"]
    note.font.name = "Arial Unicode MS"
    set_font_family(note._element.rPr.rFonts)
    note.font.size = Pt(9)
    note.font.color.rgb = RGBColor.from_string(MUTED)
    note.paragraph_format.space_after = Pt(4)
    note.paragraph_format.line_spacing = 1.12


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    add_text(p, "FRONTMIND  |  深圳鲲鹏青年创新创业项目", bold=True, color=MUTED, size=8.5)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    add_text(p, "项目申请简介（草拟版）  ·  ")
    add_field(p.add_run(), "PAGE")


def add_cover(doc: Document) -> None:
    banner = doc.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(banner)
    set_table_borders(banner, color=PURPLE, size=0)
    cell = banner.cell(0, 0)
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, top=260, start=220, bottom=260, end=220)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    add_text(p, "SHENZHEN KUNPENG YOUTH INNOVATION & ENTREPRENEURSHIP", bold=True, color=WHITE, size=9)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    add_text(p, "FrontMind", bold=True, color=PURPLE, size=17)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    add_text(p, "可信智能体驱动的企业 GEO 与 AI 增长平台", bold=True, color=NAVY, size=25)

    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    add_text(p, "深圳鲲鹏青年创新创业项目 · 项目申请简介", color=MUTED, size=13)

    accent = doc.add_table(rows=1, cols=3)
    accent.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(accent, 9360, 120)
    set_table_borders(accent, color=WHITE, size=0)
    for cell, fill, width in zip(accent.rows[0].cells, (PURPLE, GOLD, PALE_BLUE), (3560, 2240, 3560)):
        set_cell_width(cell, width)
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=35, start=0, bottom=35, end=0)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    meta_rows = [
        ["申报主体（暂定）", "深圳市超前无限科技有限公司（最终以营业执照为准）"],
        ["项目负责人", "夏凡增"],
        ["所属领域", "人工智能（主赛道）；软件和信息服务（关联支撑）"],
        ["建议推荐单位", "香港中文大学（深圳）（待校方书面确认）"],
        ["文件版本", "2026 年 7 月 · 草拟版 / 非提交版"],
        ["填报说明", "黄色标识为提交前待核验内容；公开信息须以合同、证照或权属文件举证。"],
    ]
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 8200, 700)
    set_table_borders(table, color=LINE, size=5)
    for idx, (label, value) in enumerate(meta_rows):
        row = table.add_row()
        keep_row_together(row)
        for cell, width in zip(row.cells, (2100, 6100)):
            set_cell_width(cell, width)
            set_cell_margins(cell, top=110, start=150, bottom=110, end=150)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], PALE_BLUE)
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        add_text(p0, label, bold=True, color=NAVY, size=9.5)
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        highlight = PALE_YELLOW if idx in (0, 3, 5) else None
        add_text(p1, value, size=9.5, highlight=highlight)

    doc.add_page_break()


def add_executive_summary(doc: Document) -> None:
    add_heading(doc, "申报摘要", 1)
    add_callout(
        doc,
        "一句话定位",
        "FrontMind 将企业分散的品牌事实、产品证据与市场需求，转化为可被生成式 AI 准确理解、可信引用并持续优化的知识与增长闭环。",
        fill=PURPLE_LIGHT,
        title_color=PURPLE,
    )
    add_body(
        doc,
        "FrontMind 起源于香港中文大学（深圳）数据科学学院智能决策实验室的研究与成果转化实践，面向生成式人工智能重塑“发现—比较—决策—推荐”链路的新环境，建设可信智能体驱动的企业 GEO（Generative Engine Optimization，生成式引擎优化）与 AI 增长平台。项目以企业知识与证据治理为底座，通过跨平台监测、品牌语义诊断、问答与内容资产生成、增长线索触达和企业级工作流集成，帮助企业降低 AI 回答中的事实缺失、引用不足、定位偏差与执行断层。",
    )
    add_body(
        doc,
        "项目主匹配深圳“20+8”产业集群中的人工智能产业集群，并与软件和信息服务产业集群形成交叉支撑；服务对象可覆盖深圳先进制造、科技服务、消费品牌和出海企业。现阶段已形成公开官网、企业服务平台、六类主流 AI 平台的监测框架、知识库与问答体系、内容与渠道交付模块，以及合同、支付、账号开通和权限审计等产品化基础。商业模式采用“标准化 SaaS/订阅 + 专家交付 + 企业 FDE/定制部署”的分层结构。",
    )
    add_body(
        doc,
        "项目负责人夏凡增具备计算机、自然语言处理、云计算与生成式 AI 决策研究经历，团队依托实验室在可信 AI、在线决策、强化学习和 GEO 等方向的研究积累推进成果转化。项目拟以深圳为产品研发、企业交付与市场拓展中心，优先完成证据型知识引擎、跨模型评测、智能体决策与企业安全部署的产品化，形成可重复销售和可持续续费的造血能力。",
    )
    add_callout(
        doc,
        "申报口径提示",
        "本稿中的“100+ 行业客户/品牌案例、实验室 52 篇论文与 4 项专利、获奖与合作经历”均按公开资料表述。提交前须补充合同、授权、获奖证书及知识产权权属证明；实验室成果不得在无许可或转让文件的情况下表述为企业自有资产。",
        fill=PALE_YELLOW,
    )

    add_heading(doc, "事实状态图例", 2)
    add_table(
        doc,
        ["标识", "含义", "本稿处理方式"],
        [
            ["已核实", "来自政策原文、公开官网或项目代码", "可作为事实描述，仍建议保留原始材料"],
            ["公开披露", "来自项目方/实验室公开页面", "以“公开资料显示”表述，经营业绩另行举证"],
            ["规划目标", "团队拟定的未来经营与就业目标", "不得与已实现收入、订单混写"],
            ["待补证据", "证照、身份、财务、订单、融资、IP 等未提供", "黄色标识，提交前必须替换或删除"],
        ],
        [1300, 2780, 5280],
        first_col_fill=PALE_BLUE,
        font_size=9,
    )


def add_dimension_one(doc: Document) -> None:
    add_heading(doc, "一、产业导向与商业模式", 1)
    add_heading(doc, "1.1 与深圳重点产业的匹配", 2)
    add_table(
        doc,
        ["匹配层级", "产业方向", "项目对应能力", "落地价值"],
        [
            [
                "主赛道",
                "人工智能产业集群",
                "可信智能体、生成式 AI 评测、企业知识治理、跨模型监测与决策优化",
                "形成面向企业的 AI 软件产品与专业服务，促进人工智能规模化应用",
            ],
            [
                "关联赛道",
                "软件和信息服务产业集群",
                "多租户平台、权限审计、API/工作流集成、支付与服务交付系统",
                "推动研究方法工程化、软件化和可持续交付",
            ],
            [
                "应用牵引",
                "先进制造、科技服务、消费与出海",
                "产品事实、技术证据、行业问答和多语种品牌语义资产",
                "提升深圳企业在全球 AI 搜索和智能决策入口中的可信可见度",
            ],
        ],
        [1200, 1700, 3300, 3160],
        first_col_fill=PALE_BLUE,
        font_size=8.8,
    )
    add_body(
        doc,
        "深圳“20+8”产业体系将人工智能、软件和信息服务列入重点产业集群；深圳人工智能政策同时强调基础数据、软件产品、行业应用与生态建设。FrontMind 并非单一营销工具，而是将可信 AI 方法、企业知识工程和增长场景结合，面向重点产业企业提供“可验证、可评估、可执行”的 AI 原生增长基础设施。",
    )

    add_heading(doc, "1.2 行业痛点与市场需求", 2)
    for item in (
        "企业的品牌、产品、技术与合规信息分散在官网、文档、媒体和员工经验中，生成式 AI 难以稳定读取并形成一致认知。",
        "传统 SEO 主要优化网页排名，无法充分解释大模型回答中的是否提及、排序、情绪、引用来源和竞争品牌占位。",
        "单次提问存在模型波动；缺少跨平台、重复采样和可追溯证据，导致企业难以判断改善是否真实发生。",
        "监测、内容生产、渠道分发和销售线索彼此割裂，诊断结果不能形成可持续执行与复测闭环。",
        "企业部署 AI 还需解决权限、数据安全、知识更新、审计与现有系统集成，通用工具难以直接进入生产流程。",
    ):
        add_bullet(doc, item)

    add_heading(doc, "1.3 产品与服务形态", 2)
    add_table(
        doc,
        ["产品层", "核心能力", "主要交付物", "客户价值"],
        [
            [
                "MindPromise\n可信品牌认知",
                "品牌事实图谱、语义资产审计、问题组合、六平台重复监测、引用/情绪/竞品 Gap 分析",
                "知识库、诊断报告、问答树、证据化内容资产、持续监测看板",
                "让 AI 更准确地理解、引用和推荐企业",
            ],
            [
                "MindReach\n意图与增长",
                "需求信号识别、回答逻辑设计、内容与渠道编排、线索触达和效果回流",
                "机会清单、渠道任务、内容包、线索与转化跟踪",
                "把“被看见”连接到获客与增长动作",
            ],
            [
                "MindNexus / FDE\n企业部署",
                "私有知识、流程编排、API/系统集成、多租户权限、审计与安全控制",
                "企业工作台、定制智能体、部署与运维服务",
                "把 AI 能力嵌入真实业务流程并可控运行",
            ],
        ],
        [1700, 3250, 2450, 1960],
        first_col_fill=PALE_BLUE,
        font_size=8.5,
    )
    p = doc.add_paragraph(style="Small Note")
    add_text(
        p,
        "当前监测框架覆盖豆包、腾讯元宝、DeepSeek、百度 AI+、通义千问、Kimi 等平台，并采用单轮多次回答采样；平台覆盖范围将随产品和合规要求动态调整。",
        color=MUTED,
        size=9,
    )

    add_heading(doc, "1.4 商业模式与成交路径", 2)
    add_numbered(doc, "低门槛诊断获客：以单问题、多平台现状检测验证痛点，形成明确的证据差距和优先级。")
    add_numbered(doc, "标准化订阅增收：按月或按季度交付知识库、监测、问答、内容与进度报告，形成持续续费。")
    add_numbered(doc, "企业级客单提升：针对知识复杂、系统众多或合规要求高的客户，提供 FDE、定制工作流与私有化集成。")
    add_numbered(doc, "伙伴渠道扩张：与高校、园区、行业协会、服务商及产业链龙头形成联合诊断、联合交付和客户转介。")
    add_body(
        doc,
        "收入结构计划由三部分构成：标准化软件/订阅收入、策略与内容交付收入、企业级部署及联合研发收入。项目将以产品化降低重复交付成本，以统一数据模型和交付流程提高毛利与复购，并通过持续监测形成长期客户关系。",
    )


def add_dimension_two(doc: Document) -> None:
    add_heading(doc, "二、核心团队与创业实践", 1)
    add_heading(doc, "2.1 项目负责人", 2)
    add_body(
        doc,
        "夏凡增为香港中文大学（深圳）数据科学方向博士研究生，研究聚焦 GEO 与营销智能体、上下文决策及稳健大模型应用；拥有纽约大学计算机科学硕士背景，本科阶段接受计算机工程训练。公开履历显示，其曾在清华大学从事自然语言处理算法工作，并有亚马逊云科技全栈工程经历，具备从算法研究、云端工程到企业产品化的复合能力。自 2026 年起担任 FrontMind 联合创始人兼 CEO，负责产品战略、核心技术、融资与关键客户。",
    )

    add_heading(doc, "2.2 研究与成果转化基础", 2)
    add_body(
        doc,
        "项目源自香港中文大学（深圳）智能决策实验室的研究与产业实践。实验室由李同欣博士主持，公开资料显示团队成立于 2022 年，拥有 14 名研究成员，累计发表 52 篇论文、拥有 4 项专利，研究方向涵盖可信人工智能、在线决策、强化学习、控制与 GEO。团队成员具有清华大学、加州理工学院、纽约大学、柏林工业大学、上海交通大学等教育背景，以及亚马逊、谷歌、字节跳动、华为等产业经历。",
    )
    add_callout(
        doc,
        "权属边界",
        "上述论文、专利和合作首先属于实验室或相应权利主体。本项目仅将其表述为研究与成果转化基础；如需写入“企业自有知识产权”，必须提供专利权人证明、许可/转让协议或共同研发合同。",
        fill=PALE_YELLOW,
    )

    add_heading(doc, "2.3 核心岗位与拟定分工", 2)
    add_table(
        doc,
        ["岗位", "主要职责", "现阶段说明"],
        [
            ["项目负责人 / CEO", "产品战略、技术路线、融资、关键客户与资源协同", "夏凡增；身份、年龄、学历和社保材料待提交"],
            ["算法与研究", "GEO 评测、可信智能体、决策优化、模型与数据实验", "核心成员姓名、履历、劳动/合作关系待补"],
            ["平台工程", "前后端、数据架构、账号权限、支付、部署与运维", "核心成员姓名、履历、劳动/合作关系待补"],
            ["策略与内容", "品牌事实治理、问答架构、内容标准、客户交付", "核心成员姓名、履历、代表案例待补"],
            ["商务与客户成功", "行业拓展、伙伴合作、续费和交付质量管理", "核心成员姓名、履历、在手线索待补"],
            ["学术顾问 / 合作资源", "研究指导、联合课题、人才与产业资源连接", "不等同于企业员工；合作边界和授权待确认"],
        ],
        [1800, 4200, 3360],
        first_col_fill=PALE_BLUE,
        font_size=8.7,
    )
    add_body(
        doc,
        "按照管理办法，项目团队应不少于 2 人。提交前必须确定除负责人外至少 1 名真实核心成员，并补充其学历、专业、经历、分工、劳动或合作关系证明；不得以实验室公开成员名单替代申报团队。",
    )

    add_heading(doc, "2.4 创业与落地实践", 2)
    for item in (
        "公开产品体系已形成“品牌认知—增长触达—企业部署”的分层路线，并搭建可在线访问的网站与服务入口。",
        "项目代码已具备多租户权限、企业知识库、监测看板、服务订单、合同、支付、账号开通、交付进度与审计等生产化模块。",
        "实验室公开页面披露项目团队曾获得 2024 年“未来香港”创新科技大赛冠军，并在 2025 年中国创新创业大赛、上海海聚英才全球创新创业大赛、C-Star 等活动中获奖或入选。",
        "公开页面展示了与电力、ICT 等领域机构的研究合作经历，可用于说明产业问题理解与协同研发基础；是否属于申请企业的合同或订单须以协议主体核验。",
    ):
        add_bullet(doc, item)


def add_dimension_three(doc: Document) -> None:
    add_heading(doc, "三、创新能力与市场潜力", 1)
    add_heading(doc, "3.1 核心创新与技术壁垒", 2)
    add_table(
        doc,
        ["技术层", "创新点", "可沉淀资产"],
        [
            ["证据层", "将企业主张映射到来源、时效、权威度与可验证状态，区分支持、冲突、遗漏和不可核验", "事实图谱、证据链、权威来源库"],
            ["评测层", "跨模型、跨问题、多次采样，测量提及率、排序、情绪、引用和竞品占位，降低单次回答噪声", "评测数据集、指标体系、行业基线"],
            ["决策层", "结合上下文决策与稳健智能体方法，对问题、内容和渠道动作进行优先级排序", "策略规则、决策模型、反馈数据"],
            ["执行层", "把诊断直接转为问答树、内容资产、渠道任务和复测计划，形成优化闭环", "工作流、模板、行业知识包"],
            ["工程层", "多租户权限、密钥保护、审计、支付、账号开通和企业系统集成", "企业平台、部署组件、运维标准"],
        ],
        [1300, 5140, 2920],
        first_col_fill=PALE_BLUE,
        font_size=8.6,
    )
    add_body(
        doc,
        "项目壁垒并非单一提示词或内容生成能力，而是“可信证据—跨平台评测—智能决策—执行复测—企业工程”的连续系统。随着服务行业、问题组合、模型平台和结果反馈增加，项目可积累行业基线、语义资产模板、任务数据与交付知识，形成数据网络效应和切换成本。",
    )

    add_heading(doc, "3.2 自有知识产权与合规状态", 2)
    add_body(
        doc,
        "项目已形成自研平台代码、数据结构、评测流程、行业问答方法、交付模板与智能体工作流。由于尚未提供软件著作权、专利、商标、域名权属、开源许可清单及研发人员职务成果协议，本稿不将其写为已登记的企业知识产权。",
    )
    add_callout(
        doc,
        "提交前建议形成的 IP 证据包",
        "软件著作权/专利/商标清单；代码仓库和版本记录；域名与产品页面权属；核心成员保密及职务成果协议；第三方模型、数据与开源组件许可清单；实验室成果的许可、转让或共同研发文件。",
        fill=PALE_YELLOW,
    )

    add_heading(doc, "3.3 与同类方案的差异化", 2)
    add_table(
        doc,
        ["方案类型", "典型能力", "主要局限", "FrontMind 差异化"],
        [
            ["传统 SEO / 内容服务", "关键词、网页排名、内容投放", "难解释生成式回答、引用和跨模型波动", "以 AI 回答和证据链为优化对象"],
            ["AI 可见度监测工具\n（如 Profound、Scrunch、Peec）", "品牌提及、排名、情绪、引用与竞品监测", "公开产品多以监测分析为主，企业知识和本地交付深度各异", "连接知识库、问答、内容、渠道、复测与中文平台交付"],
            ["通用 AI 咨询 / 开发", "战略咨询、模型应用、定制系统", "项目制比例高，复用和持续测量不足", "以标准化平台承接诊断，以 FDE 完成企业深度集成"],
        ],
        [1900, 2400, 2520, 2540],
        first_col_fill=PALE_BLUE,
        font_size=8.3,
    )
    p = doc.add_paragraph(style="Small Note")
    add_text(
        p,
        "竞品判断基于其公开产品页的功能分类，用于说明赛道位置，不构成对具体企业性能、价格或市场份额的绝对比较。",
        color=MUTED,
        size=9,
    )

    add_heading(doc, "3.4 当前进展", 2)
    for item in (
        "已上线 FrontMind 官方网站、产品说明与在线服务入口，形成 MindPromise、MindReach、MindNexus/FDE 的产品架构。",
        "已实现六类中文主流 AI 平台监测框架，以及知识库、问题选择、意图优化、回答逻辑、持续监测、渠道分发、进度报告和内容资产模块。",
        "已建立从服务选择、合同与支付、账号开通到项目交付的业务闭环基础，具备由研究原型向商业化产品演进的工程条件。",
        "FrontMind 官网公开披露服务“100+ 行业客户/品牌案例”。该数字在申报时须以合同、订单、发票、付款记录或经授权案例清单核验。",
    ):
        add_bullet(doc, item)

    add_heading(doc, "3.5 市场空间与增长逻辑", 2)
    add_body(
        doc,
        "需求侧正在快速形成：截至 2025 年末，我国生成式人工智能用户规模达 6.02 亿、普及率 42.8%，完成备案的生成式人工智能服务超过 750 款；2025 年我国人工智能核心产业规模超过 1.2 万亿元。深圳 2025 年人工智能核心产业规模超过 2200 亿元、规模以上人工智能企业超过 2600 家，并拥有约 2.6 万家国家高新技术企业。上述企业均面临 AI 时代知识、品牌、产品和渠道重构需求。",
    )
    add_body(
        doc,
        "GEO 尚缺少统一、可直接引用的官方市场统计口径，因此本项目采用自下而上的可验证测算。以深圳约 2.6 万家国家高新技术企业为初始潜在客户池，若首阶段触达并转化 1%，按年度综合客单价 3 万元测算，对应约 780 万元年服务规模；该测算仅为经营假设，不代表已获得订单或收入。未来可向大湾区、全国科技企业与出海品牌扩展，并通过行业知识包、伙伴渠道和企业级部署提升客单价与续费率。",
    )


def add_dimension_four(doc: Document) -> None:
    add_heading(doc, "四、运营及造血能力", 1)
    add_heading(doc, "4.1 现有运营基础", 2)
    add_body(
        doc,
        "项目已具备独立品牌、公开网站、产品架构和企业服务平台，并形成从诊断、合同、支付、账号开通到持续交付的流程基础。相较仅提供研究报告的项目，FrontMind 已将知识库、监测、权限、审计、服务订单与交付进度纳入统一系统，为后续标准化销售、续费和规模化交付提供条件。",
    )

    add_heading(doc, "4.2 经营、融资与资源现状（待企业财务确认）", 2)
    add_table(
        doc,
        ["事项", "当前申报口径", "提交所需证据"],
        [
            ["企业注册", "深圳市超前无限科技有限公司（暂定）", "营业执照、注册地址、成立日期、统一社会信用代码"],
            ["营收与回款", "【待补：2025 年、2026 年截至申报月的收入与回款】", "财务报表、纳税申报、银行流水、发票"],
            ["在手订单", "【待补：合同数量、未税金额、交付周期、已回款】", "合同关键页、订单、验收与回款证明"],
            ["融资进度", "【待补：自筹、股东投入、外部融资轮次与金额】", "出资凭证、投资协议或无外部融资说明"],
            ["合作资源", "【待补：企业名下已签协议；实验室合作单独列示】", "协议主体页、盖章页、合作内容与期限"],
            ["人员与社保", "【待补：全职/兼职人数、核心成员与社保情况】", "花名册、劳动/合作协议、社保清单"],
        ],
        [1700, 3800, 3860],
        first_col_fill=PALE_BLUE,
        font_size=8.5,
    )
    add_callout(
        doc,
        "不得混写",
        "公开案例、意向客户、实验室合作和企业已签订单必须分别统计。只有合同主体为申请企业、金额与履约状态可验证的项目，方可列入在手订单或营业收入。",
        fill=PALE_YELLOW,
    )

    add_heading(doc, "4.3 三年产值与就业目标（建议值，须负责人确认）", 2)
    add_table(
        doc,
        ["年度", "付费客户目标", "综合年客单价假设", "预计新增产值", "累计新增就业"],
        [
            ["2027", "100 家", "3 万元", "300 万元", "8 人"],
            ["2028", "200 家", "4 万元", "800 万元", "18 人"],
            ["2029", "360 家", "5 万元", "1,800 万元", "35 人"],
        ],
        [1250, 1800, 2100, 2100, 2110],
        first_col_fill=PALE_GREEN,
        font_size=9,
    )
    add_body(
        doc,
        "上述目标按“付费客户数 × 综合年客单价”测算，三年预计新增产值合计 2,900 万元。增长假设来自三个杠杆：标准化监测与知识服务扩大客户数，行业知识包与伙伴渠道降低获客/交付成本，企业 FDE 与定制部署提高客单价。就业岗位主要分布于算法研究、平台工程、数据与知识运营、行业交付、商务及客户成功。",
    )

    add_heading(doc, "4.4 资金使用与自我造血路径", 2)
    add_table(
        doc,
        ["阶段", "重点投入", "可衡量里程碑"],
        [
            ["0—6 个月", "证据型知识引擎、跨平台评测、产品安全与 IP 确权", "完成核心版本、测试集、软件著作权/专利申报及首批标杆客户"],
            ["6—18 个月", "行业知识包、伙伴渠道、交付自动化与客户成功", "形成稳定续费、缩短交付周期、提高订阅收入占比"],
            ["18—36 个月", "企业 FDE、私有部署、全国与出海市场", "提升大客户占比、实现多行业复制与正向经营现金流"],
        ],
        [1650, 3750, 3960],
        first_col_fill=PALE_BLUE,
        font_size=8.7,
    )
    add_body(
        doc,
        "项目将优先以创始团队自筹、客户回款和经营现金流支持研发与交付，并根据发展阶段引入产业资本或科研合作资金。根据管理办法，推荐或邀请入库项目可获得最高 5 万元一次性创新奖励；后续符合条件并在深圳实际运营的项目，可另行申请不超过研发投入 50%、最高 100 万元的研发补助。两类支持应按不同申报条件和程序管理，不纳入已实现收入。",
    )


def add_compact_version(doc: Document) -> None:
    add_heading(doc, "五、可直接粘贴的项目简介（精简版）", 1)
    add_body(
        doc,
        "FrontMind 是源自香港中文大学（深圳）智能决策实验室研究与成果转化实践的企业 GEO 与 AI 增长平台，主匹配深圳“20+8”产业集群中的人工智能产业集群，并与软件和信息服务产业集群形成交叉支撑。面向生成式 AI 重塑企业发现、比较与决策入口的新趋势，项目解决企业品牌与产品事实分散、AI 回答易失真、跨平台效果难衡量、监测与执行割裂以及企业级部署安全等痛点。产品由 MindPromise 可信品牌认知、MindReach 意图与增长、MindNexus/FDE 企业部署三层构成，提供事实与证据治理、六平台重复监测、引用/情绪/竞品分析、问答与内容资产、渠道触达、企业工作流集成等服务，采用“标准化 SaaS/订阅 + 专家交付 + 企业级定制部署”的商业模式。",
    )
    add_body(
        doc,
        "项目负责人夏凡增具备计算机科学、自然语言处理、云计算和生成式 AI 决策研究背景，曾有高校算法研发及云计算工程实践，自 2026 年起担任 FrontMind 联合创始人兼 CEO。项目依托实验室在可信 AI、在线决策、强化学习与 GEO 方向的研究积累，已形成公开官网、企业服务平台、知识库、监测看板、问题与内容资产、合同支付、账号开通、权限与审计等产品化模块。核心创新在于以证据型知识图谱为底座，通过跨模型重复采样减少单次回答噪声，以稳健智能体决策连接问题选择、内容生产、渠道执行与复测，并以多租户权限和系统集成进入企业生产环境。与传统 SEO、单一 AI 监测工具或项目制咨询相比，项目更强调中文主流平台、证据可信度和端到端执行闭环。",
    )
    add_body(
        doc,
        "公开资料显示，相关团队拥有多学科研究与产业经验，曾在创新创业赛事中获奖，FrontMind 官网披露服务 100+ 行业客户/品牌案例；以上经营与合作信息将在申报时以合同、订单、授权和证书核验。市场方面，截至 2025 年末我国生成式 AI 用户达 6.02 亿，深圳人工智能核心产业规模超过 2200 亿元，企业对 AI 原生知识、品牌和增长基础设施的需求快速增长。项目拟以深圳为研发和交付中心，建议三年新增产值目标分别为 300 万元、800 万元和 1800 万元，累计新增就业 35 人，重点吸纳算法、工程、数据、交付和客户成功人才。企业实际营收、融资、在手订单、自筹资金及合作协议将以财务与法律文件补充，规划目标不作为已实现业绩。",
    )


def add_submission_checklist(doc: Document) -> None:
    add_heading(doc, "六、提交前核验清单", 1)
    checklist = [
        "企业：营业执照、统一社会信用代码、注册地址、成立日期、股权结构，确认申请主体法定全称。",
        "负责人：身份证号码、出生日期、手机号、最高学历及证明；核验年龄和学历是否符合管理办法。",
        "团队：至少 2 人的真实核心名单、专业背景、分工、劳动/合作关系及社保材料。",
        "知识产权：专利、软著、商标、域名、代码与数据权属；实验室成果另附许可、转让或共同研发文件。",
        "经营：2025 年及 2026 年截至申报月的财务报表、纳税、流水、发票、客户合同、订单、验收和回款。",
        "融资与自筹：股东投入、银行余额、融资协议或无外部融资说明，明确可支撑的研发周期。",
        "案例与合作：对“100+”公开披露形成可审计清单；合作协议确认主体、期限、金额和履约状态。",
        "获奖与媒体：奖项证书、主办方通知、报道链接；明确获奖主体是个人、实验室、团队还是企业。",
        "测算：由负责人确认 2027—2029 年客户数、客单价、产值、成本、现金流和新增就业目标。",
        "推荐：由香港中文大学（深圳）确认推荐资格、校内流程、盖章主体、联系人与截止日期。",
    ]
    for item in checklist:
        add_bullet(doc, item)
    add_callout(
        doc,
        "最终提交动作",
        "将黄色占位全部替换；对照政策准备诚信承诺、知识产权合规声明、专项审计及社保材料；确保推荐表、申请简介、证照、财务与合同中的企业名称、负责人、金额和时间完全一致。",
        fill=PALE_YELLOW,
    )


def add_sources(doc: Document) -> None:
    add_heading(doc, "主要依据与公开来源", 1)
    sources: list[tuple[str, list[tuple[str, str]]]] = [
        (
            "深圳市科技创新局，《深圳鲲鹏青年创新创业项目管理办法》（深科技创新规〔2025〕2 号），用户提供 PDF。",
            [],
        ),
        (
            "深圳市科技创新局，2026 年深圳鲲鹏青年创新创业项目通知。",
            [("官方通知", "https://stic.sz.gov.cn/xxgk/tzgg/content/post_12903809.html")],
        ),
        (
            "深圳市人民政府：“20+8”产业集群、人工智能政策与 2025 年产业数据。",
            [
                ("产业集群", "https://www.sz.gov.cn/cn/zjsz/gl/content/post_12553610.html"),
                ("人工智能政策", "https://www.sz.gov.cn/zfgb/2025/gb1368/content/post_12114699.html"),
                ("产业数据", "https://www.sz.gov.cn/cn/xxgk/zfxxgj/zwdt/content/post_12896926.html"),
            ],
        ),
        (
            "国家统计局与国家互联网信息办公室：人工智能产业、生成式人工智能用户与服务备案数据。",
            [
                ("国家统计局", "https://www.stats.gov.cn/sj/sjjd/202606/t20260602_1963861.html"),
                ("国家网信办", "https://www.cac.gov.cn/2026-04/16/c_1778078563279563.htm"),
            ],
        ),
        ("FrontMind 官方网站与产品页面。", [("frontmind.net", "https://www.frontmind.net/")]),
        (
            "香港中文大学（深圳）智能决策实验室官方网站、团队与新闻页面。",
            [("cuhkgeo.com", "https://www.cuhkgeo.com/")],
        ),
        ("夏凡增个人学术主页。", [("个人主页", "https://xiafanzeng.github.io/")]),
        (
            "Profound、Scrunch、Peec 公开产品页（仅用于竞品能力分类）。",
            [
                ("Profound", "https://www.tryprofound.com/features"),
                ("Scrunch", "https://scrunch.com/"),
                ("Peec", "https://peec.ai/"),
            ],
        ),
        (
            "FrontMind 项目代码仓库：产品、监测、知识库、支付、合同、账号、权限与审计等模块（本地项目文件夹）。",
            [],
        ),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 9360, 120)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for cell, header, width in zip(table.rows[0].cells, ("依据", "链接"), (6900, 2460)):
        set_cell_width(cell, width)
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=80, start=110, bottom=80, end=110)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_text(p, header, bold=True, color=WHITE, size=8.5)
    for index, (label, links) in enumerate(sources):
        row = table.add_row()
        keep_row_together(row)
        for cell, width in zip(row.cells, (6900, 2460)):
            set_cell_width(cell, width)
            set_cell_margins(cell, top=60, start=100, bottom=60, end=100)
            if index % 2 == 1:
                set_cell_shading(cell, "F7F9FC")
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        add_text(p, label, size=8)
        p2 = row.cells[1].paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.0
        if not links:
            add_text(p2, "本地材料", color=MUTED, size=8)
        else:
            for link_index, (link_label, url) in enumerate(links):
                if link_index:
                    add_text(p2, " · ", color=MUTED, size=8)
                add_hyperlink(p2, link_label, url)
    p = doc.add_paragraph(style="Small Note")
    p.paragraph_format.space_before = Pt(6)
    add_text(
        p,
        "注：本稿为基于现有材料形成的申报草案，不替代主管部门、推荐单位或专业法律/财务意见。政策、截止时间和推荐流程应以当期正式通知为准。",
        color=MUTED,
        size=9,
    )


def core_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "FrontMind 深圳鲲鹏青年创新创业项目申请简介"
    props.subject = "深圳鲲鹏青年创新创业项目申报草稿"
    props.author = "FrontMind"
    props.keywords = "FrontMind, 深圳鲲鹏青年, 人工智能, GEO, 可信智能体"
    props.comments = "基于公开资料、政策文件与项目代码形成；黄色内容提交前需核验。"


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    core_properties(doc)

    add_cover(doc)
    add_executive_summary(doc)
    add_dimension_one(doc)
    add_dimension_two(doc)
    add_dimension_three(doc)
    add_dimension_four(doc)
    add_compact_version(doc)
    add_submission_checklist(doc)
    add_sources(doc)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
