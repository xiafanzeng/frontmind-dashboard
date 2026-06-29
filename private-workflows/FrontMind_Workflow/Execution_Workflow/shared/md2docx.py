#!/usr/bin/env python3
"""
md2docx.py - FrontMind Markdown to DOCX Converter

Converts Markdown files to professionally formatted DOCX documents
with embedded images, proper heading styles, and brand formatting.

★ 全类型铁律：
  1. 绝对禁止任何表格 — 遇到 Markdown 管道表格时报错退出，不转图片不忽略。
  2. 绝对禁止 H3 小标题 — 遇到 ### 时报错退出。
  3. 绝对禁止内部元数据表头 — 正文前30行不得含"文章ID""文章类型"等内部标记。
  4. 图片必须物理嵌入 — 支持 PNG/JPEG/WebP（WebP 通过 Pillow 转 PNG 后嵌入）。
  5. Markdown 加粗语法 **text** 必须转为 Word 加粗格式，不得残留星号。

Features:
- Heading styles (H2 only; H1 prohibited in final output, H3 triggers error)
- Image embedding (PNG/JPEG direct, WebP auto-converted via Pillow)
- Markdown bold → Word bold formatting
- Code block formatting
- Blockquote styling
- Automatic image size optimization

Usage:
    python shared/md2docx.py input.md output.docx --images ./images/ --brand "BrandName"
    python shared/md2docx.py input.md output.docx --images ./images/ --brand "BrandName" --validate
"""

import argparse
import io
import os
import re
import sys
import tempfile
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

try:
    from PIL import Image
    HAS_PILLOW = True
except ImportError:
    HAS_PILLOW = False
    print("Warning: Pillow not installed. WebP support disabled.", file=sys.stderr)

# ── Style Constants ──────────────────────────────────────────────────
FONT_BODY_CN = "Microsoft YaHei"
FONT_BODY_EN = "Calibri"
FONT_SIZE_BODY = Pt(12)
FONT_SIZE_H1 = Pt(22)
FONT_SIZE_H2 = Pt(16)
COLOR_PRIMARY = RGBColor(107, 33, 168)
COLOR_BLACK = RGBColor(26, 26, 26)
COLOR_GRAY = RGBColor(100, 116, 139)
MAX_IMAGE_WIDTH = Inches(5.5)

# ── Forbidden metadata keywords (in first 30 lines) ─────────────────
METADATA_KEYWORDS = [
    "文章ID", "文章类型", "核心GEO问题", "正文字数", "配图数量",
    "标题池", "生成日期", "FrontMind 执行层", "Execution Layer",
    "资料来源与口径说明", "本文采用"
]


def setup_styles(doc):
    """Configure document styles for FrontMind branding."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_BODY_EN
    font.size = FONT_SIZE_BODY
    font.color.rgb = COLOR_BLACK

    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.5

    for level, size in [(1, FONT_SIZE_H1), (2, FONT_SIZE_H2)]:
        heading_style = doc.styles[f"Heading {level}"]
        heading_font = heading_style.font
        heading_font.name = FONT_BODY_EN
        heading_font.size = size
        heading_font.bold = True
        heading_font.color.rgb = COLOR_PRIMARY if level == 1 else COLOR_BLACK

    return doc


def convert_webp_to_png(webp_path):
    """Convert WebP image to PNG in a temp file, return path to PNG."""
    if not HAS_PILLOW:
        return None
    try:
        img = Image.open(webp_path)
        # Create temp PNG file
        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False, prefix="frontmind_img_")
        img.save(tmp.name, format="PNG")
        tmp.close()
        return tmp.name
    except Exception as e:
        print(f"Warning: Failed to convert WebP {webp_path}: {e}", file=sys.stderr)
        return None


def add_image_to_doc(doc, image_path, alt_text=""):
    """Add an image to the document with size constraints. Supports PNG/JPEG/WebP."""
    img_path = Path(image_path)
    if not img_path.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"[Image not found: {image_path}]")
        run.font.color.rgb = RGBColor(255, 0, 0)
        return False

    actual_path = str(img_path)
    temp_png = None

    # WebP → PNG conversion
    if img_path.suffix.lower() == ".webp":
        temp_png = convert_webp_to_png(actual_path)
        if temp_png is None:
            p = doc.add_paragraph()
            run = p.add_run(f"[WebP conversion failed: {image_path}. Install Pillow.]")
            run.font.color.rgb = RGBColor(255, 0, 0)
            return False
        actual_path = temp_png

    try:
        doc.add_picture(actual_path, width=MAX_IMAGE_WIDTH)
        return True
    except Exception as e:
        p = doc.add_paragraph()
        run = p.add_run(f"[Image error: {e}]")
        run.font.color.rgb = RGBColor(255, 0, 0)
        return False
    finally:
        # Clean up temp file
        if temp_png and os.path.exists(temp_png):
            try:
                os.unlink(temp_png)
            except OSError:
                pass


def add_paragraph_with_bold(doc, text):
    """
    Add a paragraph that correctly handles Markdown bold syntax.
    **text** → Word bold formatting (no asterisks in output).
    """
    p = doc.add_paragraph()
    # Split by bold markers
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            # Bold text
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            # Also handle single * italic
            italic_parts = re.split(r'(\*[^*]+\*)', part)
            for ip in italic_parts:
                if ip.startswith("*") and ip.endswith("*") and len(ip) > 2:
                    run = p.add_run(ip[1:-1])
                    run.italic = True
                else:
                    if ip:
                        p.add_run(ip)
    return p


def parse_markdown_blocks(md_content):
    """
    Parse Markdown into structured blocks.
    
    ★ FATAL ERRORS:
    - If H3 (###) is found → raises ValueError
    - If Markdown table (|...|) is found → raises ValueError
    """
    blocks = []
    lines = md_content.split("\n")
    i = 0
    in_code = False
    code_buf = []

    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            if in_code:
                blocks.append({"type": "code", "content": "\n".join(code_buf)})
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ★ FATAL: H3 detected
        if line.startswith("### "):
            raise ValueError(
                f"[md2docx FATAL] H3 小标题被禁止！第 {i+1} 行: '{line.strip()}'\n"
                f"请退回 E2 重写，全文只允许 H2 大标题。"
            )

        # ★ FATAL: Table detected
        if line.startswith("|") and i + 1 < len(lines) and ("---" in lines[i + 1] or ":---" in lines[i + 1]):
            raise ValueError(
                f"[md2docx FATAL] Markdown 表格被禁止！第 {i+1} 行: '{line.strip()}'\n"
                f"请退回 E2 重写，正文不允许任何表格，改用序号列表描述。"
            )

        if line.startswith("# "):
            # H1 is skipped (正本不带文章标题)
            print(f"Warning: H1 found at line {i+1}, skipping (正本不带标题).", file=sys.stderr)
        elif line.startswith("## "):
            blocks.append({"type": "h2", "content": line[3:].strip()})
        elif line.startswith("> "):
            blocks.append({"type": "quote", "content": line[2:].strip()})
        elif re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line):
            m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line)
            blocks.append({"type": "image", "alt": m.group(1), "src": m.group(2)})
        elif line.startswith("---"):
            # Horizontal rule → skip
            pass
        elif line.strip():
            blocks.append({"type": "paragraph", "content": line.strip()})

        i += 1

    return blocks


def check_metadata_in_content(md_content):
    """Check first 30 lines for forbidden internal metadata."""
    lines = md_content.split("\n")[:30]
    found = []
    for i, line in enumerate(lines):
        for kw in METADATA_KEYWORDS:
            if kw in line:
                found.append(f"Line {i+1}: found '{kw}' in '{line.strip()}'")
    return found


def convert_md_to_docx(md_path, output_path, images_dir=None, brand_name=""):
    """
    Convert Markdown file to DOCX with FrontMind formatting.
    
    Raises ValueError if:
    - H3 headings are found
    - Markdown tables are found
    - Internal metadata is found in first 30 lines
    """
    with open(md_path, "r", encoding="utf-8") as f:
        md_content = f.read()

    # Strip YAML frontmatter if present
    if md_content.startswith("---"):
        end = md_content.find("---", 3)
        if end != -1:
            md_content = md_content[end + 3:].strip()

    # Check for metadata
    metadata_issues = check_metadata_in_content(md_content)
    if metadata_issues:
        raise ValueError(
            f"[md2docx FATAL] 正文中发现内部元数据（前30行）:\n" +
            "\n".join(metadata_issues) +
            "\n请退回 E2/E4 删除所有内部元数据后重试。"
        )

    # Parse blocks (will raise ValueError on H3 or tables)
    blocks = parse_markdown_blocks(md_content)

    # Create document
    doc = Document()
    setup_styles(doc)

    image_count = 0

    for block in blocks:
        btype = block["type"]

        if btype == "h2":
            doc.add_heading(block["content"], level=2)
        elif btype == "paragraph":
            text = block["content"]
            add_paragraph_with_bold(doc, text)
        elif btype == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(block["content"])
            run.font.italic = True
            run.font.color.rgb = COLOR_GRAY
        elif btype == "code":
            p = doc.add_paragraph()
            run = p.add_run(block["content"])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif btype == "image":
            src = block["src"]
            if images_dir:
                src = os.path.join(images_dir, os.path.basename(src))
            if add_image_to_doc(doc, src, block.get("alt", "")):
                image_count += 1

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    file_size = Path(output_path).stat().st_size
    print(f"Generated: {output_path}")
    print(f"  Size: {file_size / 1024:.1f} KB")
    print(f"  Images embedded: {image_count}")

    return {"path": output_path, "size": file_size, "images": image_count}


def validate_docx(docx_path):
    """Validate DOCX output meets FrontMind standards."""
    errors = []
    path = Path(docx_path)

    if not path.exists():
        return ["File not found"]

    size = path.stat().st_size
    if size < 10 * 1024:
        errors.append(f"File too small: {size / 1024:.1f} KB (expected > 10 KB with images)")
    if size > 10 * 1024 * 1024:
        errors.append(f"File too large: {size / (1024*1024):.1f} MB (max 10 MB)")

    doc = Document(docx_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)

    # Check for Markdown syntax residue
    forbidden = ["IMAGE_SLOT", "[Image placeholder]", "**", "---", "<!--", "###"]
    for f in forbidden:
        if f in full_text:
            errors.append(f"Forbidden content found: '{f}'")

    # Check for internal metadata
    first_30_lines = "\n".join(p.text for p in doc.paragraphs[:30])
    for kw in METADATA_KEYWORDS:
        if kw in first_30_lines:
            errors.append(f"Internal metadata found: '{kw}'")

    # Verify no native tables exist
    if len(doc.tables) > 0:
        errors.append(f"Native Word tables found: {len(doc.tables)} (must be 0)")

    # Verify images are embedded
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    image_parts = [rel for rel in doc.part.rels.values() if "image" in rel.reltype]
    if len(image_parts) == 0:
        errors.append("No embedded images found (expected ≥ 1)")

    return errors


def main():
    parser = argparse.ArgumentParser(description="FrontMind MD to DOCX Converter")
    parser.add_argument("input", help="Input Markdown file")
    parser.add_argument("output", help="Output DOCX file")
    parser.add_argument("--images", default="./images/", help="Images directory")
    parser.add_argument("--brand", required=True, help="Brand name")
    parser.add_argument("--validate", action="store_true", help="Validate output")
    args = parser.parse_args()

    if not Path(args.input).exists():
        print(f"Error: Input file not found: {args.input}")
        sys.exit(1)

    try:
        result = convert_md_to_docx(args.input, args.output, args.images, args.brand)
    except ValueError as e:
        print(f"\n{'='*60}\n{e}\n{'='*60}", file=sys.stderr)
        sys.exit(2)

    if args.validate:
        errors = validate_docx(args.output)
        if errors:
            print("Validation FAILED:")
            for e in errors:
                print(f"  [X] {e}")
            sys.exit(1)
        else:
            print("Validation PASSED")


if __name__ == "__main__":
    main()
