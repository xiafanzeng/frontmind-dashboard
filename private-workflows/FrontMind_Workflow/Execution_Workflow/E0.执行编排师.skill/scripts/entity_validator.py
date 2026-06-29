#!/usr/bin/env python3
"""
FrontMind 两层实体验证器 (Entity Validator)

功能：
  E0 执行编排师在 E4 交付后亲自运行此脚本，执行两层实体验证：
  - 第一层：文件实体验证（文件存在性、大小、字数）
  - 第二层：DOCX 图片嵌入验证（图片数量、格式、占位符残留、Markdown 残留）

用法：
  python3 entity_validator.py --docx path/to/article.docx --md path/to/article.md --images-dir images/ --article-id A1-001

依赖：
  pip install python-docx
"""

import argparse
import os
import re
import sys
from typing import List, Tuple

# ─── 尝试导入 python-docx ───────────────────────────────────────────
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("[WARN] python-docx 未安装，第二层验证将跳过。请运行: pip install python-docx")


# ─── 常量配置 ────────────────────────────────────────────────────────
DOCX_MIN_SIZE_KB = 100       # DOCX 最小文件大小（KB）
DOCX_MAX_SIZE_MB = 10        # DOCX 最大文件大小（MB）
MD_MIN_CHARS_A = 3500        # A 类文章最低字符数
MD_MIN_CHARS_B1 = 8000       # B1 白皮书最低字符数
MD_MIN_CHARS_C1B = 3500      # C1b 品牌新闻稿最低字符数
MD_MIN_CHARS_C1A = 800       # C1a 事件新闻稿最低字符数
IMG_MIN_SIZE_KB = 10         # 图片最小文件大小（KB）
IMG_MIN_COUNT = 2            # 最少配图数量
DOCX_IMG_MIN_COUNT = 3       # DOCX 内嵌图片最少数量

# 占位符关键词列表
PLACEHOLDER_PATTERNS = [
    "[图片占位]", "[请插入图片]", "[Python生成]", "[网络搜索下载]",
    "[AI生成]", "[待补充]", "[此处插入]", "图片占位符",
    "IMAGE_SLOT", "IMAGE-SLOT", "[图", "<!-- IMAGE",
    "[此处填写]", "TODO", "[PLACEHOLDER]"
]

# Markdown 残留检测正则
MD_RESIDUE_PATTERNS = [
    (r"\*\*[^*]+\*\*", "加粗语法 **text**"),
    (r"^#{1,6}\s", "标题语法 # heading"),
    (r"\[.*?\]\(.*?\)", "链接语法 [text](url)"),
    (r"^\|.*\|$", "表格语法 |col|col|"),
    (r"^[-*_]{3,}$", "分隔线 ---"),
    (r"^>\s", "引用块 > text"),
    (r"`[^`]+`", "行内代码 `code`"),
    (r"^```", "代码块 ```"),
    (r"<!--.*?-->", "HTML 注释 <!-- -->"),
]


class ValidationResult:
    """验证结果收集器。"""

    def __init__(self):
        self.checks: List[Tuple[bool, str]] = []

    def add(self, passed: bool, message: str) -> None:
        """添加一条检查结果。"""
        self.checks.append((passed, message))
        icon = "✅" if passed else "❌"
        print(f"  {icon} {message}")

    @property
    def all_passed(self) -> bool:
        """是否所有检查都通过。"""
        return all(passed for passed, _ in self.checks)

    @property
    def failed_count(self) -> int:
        """失败的检查数量。"""
        return sum(1 for passed, _ in self.checks if not passed)

    def summary(self) -> str:
        """生成验证摘要。"""
        total = len(self.checks)
        passed = sum(1 for p, _ in self.checks if p)
        failed = total - passed
        status = "✅ 全部通过" if self.all_passed else f"❌ {failed}/{total} 项未通过"
        return f"验证结果: {status}"


def get_min_chars(article_id: str) -> int:
    """根据文章编号推断最低字符数要求。"""
    aid = article_id.upper()
    if aid.startswith("B1"):
        return MD_MIN_CHARS_B1
    elif aid.startswith("C1A"):
        return MD_MIN_CHARS_C1A
    elif aid.startswith("C1B"):
        return MD_MIN_CHARS_C1B
    elif aid.startswith(("A", "B", "C1B")):
        return MD_MIN_CHARS_A
    else:
        return MD_MIN_CHARS_A  # 默认使用 A 类标准


def validate_layer1(docx_path: str, md_path: str, images_dir: str,
                    article_id: str) -> ValidationResult:
    """
    第一层：文件实体验证。
    检查文件存在性、大小和字数。
    """
    result = ValidationResult()
    print("\n═══ 第一层：文件实体验证 ═══")

    # 1. 检查 DOCX 文件
    if not os.path.exists(docx_path):
        result.add(False, f"DOCX 文件不存在: {docx_path}")
    else:
        size_bytes = os.path.getsize(docx_path)
        size_kb = size_bytes / 1024
        size_mb = size_bytes / (1024 * 1024)

        if size_kb < DOCX_MIN_SIZE_KB:
            result.add(False,
                       f"DOCX 文件过小: {size_kb:.0f}KB（最低 {DOCX_MIN_SIZE_KB}KB，疑似图片未嵌入）")
        elif size_mb > DOCX_MAX_SIZE_MB:
            result.add(False,
                       f"DOCX 文件过大: {size_mb:.1f}MB（超过 {DOCX_MAX_SIZE_MB}MB 限制）")
        else:
            result.add(True, f"DOCX 文件大小: {size_kb:.0f}KB")

    # 2. 检查 MD 文件
    if not os.path.exists(md_path):
        result.add(False, f"MD 文件不存在: {md_path}")
    else:
        with open(md_path, "r", encoding="utf-8") as f:
            content = f.read()
        char_count = len(content)
        min_chars = get_min_chars(article_id)

        if char_count < min_chars:
            result.add(False,
                       f"MD 文件字数不足: {char_count} 字符（{article_id} 类型要求 ≥{min_chars}）")
        else:
            result.add(True, f"MD 文件字数: {char_count} 字符")

    # 3. 检查配图文件
    if not os.path.isdir(images_dir):
        result.add(False, f"图片目录不存在: {images_dir}")
    else:
        img_count = 0
        for fname in sorted(os.listdir(images_dir)):
            if not fname.startswith(article_id):
                continue
            if not fname.lower().endswith((".png", ".jpg", ".jpeg", ".webp")):
                continue

            fpath = os.path.join(images_dir, fname)
            fsize_kb = os.path.getsize(fpath) / 1024

            if fsize_kb < IMG_MIN_SIZE_KB:
                result.add(False,
                           f"图片过小: {fname} ({fsize_kb:.1f}KB，疑似损坏)")
            else:
                result.add(True, f"图片: {fname} ({fsize_kb:.1f}KB)")
            img_count += 1

        if img_count < IMG_MIN_COUNT:
            result.add(False,
                       f"配图数量不足: {img_count} 张（最低 {IMG_MIN_COUNT} 张）")
        else:
            result.add(True, f"配图数量: {img_count} 张")

    return result


def validate_layer2(docx_path: str) -> ValidationResult:
    """
    第二层：DOCX 图片嵌入验证。
    深入检查 DOCX 内部结构：图片数量、格式、占位符残留、Markdown 残留。
    """
    result = ValidationResult()
    print("\n═══ 第二层：DOCX 图片嵌入验证 ═══")

    if not HAS_DOCX:
        result.add(False, "python-docx 未安装，无法执行第二层验证")
        return result

    if not os.path.exists(docx_path):
        result.add(False, f"DOCX 文件不存在: {docx_path}")
        return result

    doc = Document(docx_path)
    file_size = os.path.getsize(docx_path)

    # 1. 检查嵌入图片数量
    img_rels = [rel for rel in doc.part.rels.values() if "image" in rel.reltype]
    img_count = len(img_rels)

    if img_count < DOCX_IMG_MIN_COUNT:
        result.add(False,
                   f"DOCX 嵌入图片不足: {img_count} 张（最低 {DOCX_IMG_MIN_COUNT} 张）")
    else:
        result.add(True, f"DOCX 嵌入图片: {img_count} 张")

    # 2. 检查图片格式（应为 WebP）
    non_webp = []
    for rel in img_rels:
        target = rel.target_ref if hasattr(rel, "target_ref") else str(rel._target)
        if not target.lower().endswith(".webp"):
            non_webp.append(target)

    if non_webp:
        result.add(False,
                   f"发现 {len(non_webp)} 张非 WebP 格式图片: {non_webp[:3]}")
    else:
        result.add(True, "所有嵌入图片均为 WebP 格式")

    # 3. 提取全文文本
    full_text = "\n".join(p.text for p in doc.paragraphs)

    # 4. 检查占位符残留
    found_placeholders = []
    for pattern in PLACEHOLDER_PATTERNS:
        if pattern in full_text:
            found_placeholders.append(pattern)

    if found_placeholders:
        result.add(False,
                   f"发现占位符残留: {found_placeholders}")
    else:
        result.add(True, "无占位符残留")

    # 5. 检查 Markdown 语法残留
    md_residues = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        for pattern, desc in MD_RESIDUE_PATTERNS:
            if re.search(pattern, text, re.MULTILINE):
                md_residues.append(f"段落{i+1}: {desc} → {text[:60]}")
                break  # 每段只报告第一个问题

    if md_residues:
        result.add(False,
                   f"Markdown 语法残留 {len(md_residues)} 处")
        for r in md_residues[:10]:
            print(f"      {r}")
    else:
        result.add(True, "无 Markdown 语法残留")

    # 6. 检查文件大小
    max_bytes = DOCX_MAX_SIZE_MB * 1024 * 1024
    if file_size > max_bytes:
        result.add(False,
                   f"DOCX 文件超过 {DOCX_MAX_SIZE_MB}MB: {file_size/1024/1024:.1f}MB")
    else:
        result.add(True,
                   f"DOCX 文件大小: {file_size/1024:.0f}KB（≤{DOCX_MAX_SIZE_MB}MB）")

    return result


def main():
    """CLI 入口：执行两层实体验证。"""
    parser = argparse.ArgumentParser(
        description="FrontMind 两层实体验证器 — E0 在 E4 交付后亲自执行"
    )
    parser.add_argument("--docx", required=True, help="DOCX 文件路径")
    parser.add_argument("--md", required=True, help="MD 文件路径")
    parser.add_argument("--images-dir", required=True, help="图片目录路径")
    parser.add_argument("--article-id", required=True, help="文章编号（如 A1-001）")
    args = parser.parse_args()

    print(f"╔══════════════════════════════════════════════════╗")
    print(f"║  FrontMind 两层实体验证器                        ║")
    print(f"║  文章: {args.article_id:<40s}  ║")
    print(f"╚══════════════════════════════════════════════════╝")

    # 第一层验证
    r1 = validate_layer1(args.docx, args.md, args.images_dir, args.article_id)

    # 第二层验证
    r2 = validate_layer2(args.docx)

    # 综合判定
    print("\n" + "═" * 50)
    print(f"第一层 {r1.summary()}")
    print(f"第二层 {r2.summary()}")

    all_passed = r1.all_passed and r2.all_passed
    total_failed = r1.failed_count + r2.failed_count

    if all_passed:
        print("\n🎉 两层实体验证全部通过！可以进入成品展示步骤。")
        sys.exit(0)
    else:
        print(f"\n⚠️ 验证未通过（{total_failed} 项失败），需打回 E4 修复。")
        sys.exit(1)


if __name__ == "__main__":
    main()
