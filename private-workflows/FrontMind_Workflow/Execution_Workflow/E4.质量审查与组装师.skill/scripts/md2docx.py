#!/usr/bin/env python3
"""
FrontMind E4 MD→DOCX 组装器 (Markdown to DOCX Assembler)

★ 重要变更：表格以图片形式嵌入（非 Word 原生表格），
  确保最终 DOCX 仅包含纯文本+图片，兼容各类媒体平台投稿要求。
  很多媒体不接受表格形式只有纯文本和图片，原生表格会导致乱码。

功能：
  - 解析 Markdown 文件（标题/段落/列表/表格/图片）
  - 转换为 Word 格式（所有 Markdown 语法必须转为 Word 格式）
  - 表格 → 图片渲染（matplotlib 生成高质量表格图片后嵌入）
  - 图片 WebP 转换与物理嵌入
  - 文件大小控制（quality 自动递减）
  - Markdown 残留检测与清理
  - 占位符扫描

用法：
  python3 md2docx.py \\
    --input "{brand}_{article_id}_final.md" \\
    --images-dir "./images/" \\
    --output "{brand}_{article_id}_final.docx" \\
    --max-size 10

依赖：
  pip3 install python-docx Pillow matplotlib
"""

import argparse
import io
import os
import re
import sys
import tempfile
from typing import Dict, List, Optional, Tuple

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("警告: python-docx 未安装，请运行 pip3 install python-docx", file=sys.stderr)

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False
    print("警告: Pillow 未安装，图片处理功能不可用", file=sys.stderr)

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.font_manager as fm
    HAS_MPL = True
except ImportError:
    HAS_MPL = False
    print("警告: matplotlib 未安装，表格将以纯文本形式嵌入", file=sys.stderr)


# ============================================================
# 排版常量（严格遵循 E4 SKILL.md 规范）
# ============================================================

STYLE_CONFIG = {
    'h1': {'size': Pt(22), 'bold': True, 'color': RGBColor(0x1A, 0x1A, 0x1A)},
    'h2': {'size': Pt(16), 'bold': True, 'color': RGBColor(0x2C, 0x3E, 0x50)},
    'h3': {'size': Pt(13), 'bold': True, 'color': RGBColor(0x34, 0x49, 0x5E)},
    'body': {'size': Pt(11), 'bold': False, 'color': RGBColor(0x33, 0x33, 0x33)},
    'caption': {'size': Pt(9), 'bold': False, 'color': RGBColor(0x66, 0x66, 0x66)},
}

LINE_SPACING = 1.5
PAGE_WIDTH_INCHES = 6.0  # A4 页面可用宽度（减去边距）

# ── Table Image Rendering Constants ───────────────────────────────────────
# 图片宽度固定为 Word 页面可用宽度 (15.5cm ≈ 6.1 inch)
# 插入 Word 时以 15.5cm 宽度擑满页面，确保文字不被缩小
TABLE_FIG_WIDTH_INCH = 6.1
TABLE_DPI = 200
TABLE_ROW_HEIGHT = 0.38  # 紧凑行高，减少留白
TABLE_HEADER_HEIGHT = 0.42
TABLE_HEADER_BG = "#3B3B5C"
TABLE_HEADER_TEXT = "#FFFFFF"
TABLE_ROW_BG_EVEN = "#F8F8FA"
TABLE_ROW_BG_ODD = "#FFFFFF"
TABLE_BORDER_COLOR = "#E0E0E0"
TABLE_TEXT_COLOR = "#2D2D2D"
TABLE_FONT_SIZE = 10.5  # 与正文小四(12pt)视觉一致
TABLE_HEADER_FONT_SIZE = 11
TABLE_INSERT_WIDTH_CM = 15.5  # 插入 Word 时的图片宽度(cm)

# 中文字体路径（优先使用 Noto Sans CJK SC）
_CJK_FONT_PATHS = [
    "/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
    "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
]
_CJK_BOLD_FONT_PATHS = [
    "/usr/share/fonts/opentype/noto/NotoSansCJKsc-Bold.otf",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc",
]

# Markdown 残留检测模式
MARKDOWN_PATTERNS = [
    (r'\*\*([^*]+)\*\*', 'bold_markdown', r'\1'),
    (r'__([^_]+)__', 'bold_markdown_alt', r'\1'),
    (r'\*([^*]+)\*', 'italic_markdown', r'\1'),
    (r'_([^_]+)_', 'italic_markdown_alt', r'\1'),
    (r'~~([^~]+)~~', 'strikethrough_markdown', r'\1'),
    (r'`([^`]+)`', 'code_markdown', r'\1'),
    (r'\[([^\]]+)\]\(([^\)]+)\)', 'link_markdown', r'\1'),
]

# 占位符检测模式
PLACEHOLDER_PATTERNS = [
    r'IMAGE_SLOT[-:]?\s*\w+',
    r'\{\{[^}]+\}\}',
    r'\[TODO\]',
    r'\[TBD\]',
    r'\[待定\]',
    r'\[此处填写\]',
    r'\[待补充\]',
    r'caption:',
]


# ============================================================
# 表格图片渲染引擎
# ============================================================

def _get_cjk_font(bold=False):
    """获取 CJK 中文字体用于 matplotlib 渲染。"""
    paths = _CJK_BOLD_FONT_PATHS if bold else _CJK_FONT_PATHS
    for fp in paths:
        if os.path.exists(fp):
            return fm.FontProperties(fname=fp)
    return fm.FontProperties(family="sans-serif")


def render_table_as_image(headers: List[str], rows: List[List[str]], dpi: int = None) -> Optional[bytes]:
    """
    将表格渲染为 PNG 图片。
    图片宽度固定为 Word 页面可用宽度，插入后擑满页面，
    确保文字与正文字号视觉一致。

    使用 matplotlib 生成高质量表格图片，支持中文字体，
    确保在各类媒体平台上不会出现乱码问题。

    Args:
        headers: 表头列表
        rows: 数据行列表
        dpi: 图片 DPI（默认使用 TABLE_DPI）

    Returns:
        PNG 图片字节数据，渲染失败返回 None
    """
    if not HAS_MPL:
        return None

    if dpi is None:
        dpi = TABLE_DPI

    num_cols = len(headers)
    num_rows = len(rows)

    # 计算列宽比例（基于内容长度）
    col_weights = []
    for j in range(num_cols):
        max_len = sum(2 if ord(c) > 127 else 1 for c in str(headers[j]))
        for row in rows:
            if j < len(row):
                cell_len = sum(2 if ord(c) > 127 else 1 for c in str(row[j]))
                max_len = max(max_len, cell_len)
        col_weights.append(max(max_len, 2))

    total_weight = sum(col_weights)
    col_widths = [w / total_weight * TABLE_FIG_WIDTH_INCH for w in col_weights]

    # 图片尺寸 - 紧凑布局，最小留白
    fig_height = TABLE_HEADER_HEIGHT + num_rows * TABLE_ROW_HEIGHT + 0.05
    fig_width = TABLE_FIG_WIDTH_INCH

    fig, ax = plt.subplots(figsize=(fig_width, fig_height), dpi=dpi)
    ax.set_xlim(0, fig_width)
    ax.set_ylim(0, TABLE_HEADER_HEIGHT + num_rows * TABLE_ROW_HEIGHT)
    ax.axis("off")
    fig.subplots_adjust(left=0, right=1, top=1, bottom=0)

    font_prop = _get_cjk_font(bold=False)
    font_prop_bold = _get_cjk_font(bold=True)

    table_height = TABLE_HEADER_HEIGHT + num_rows * TABLE_ROW_HEIGHT
    y_top = table_height

    # 绘制表头行
    x = 0
    for col_idx, header in enumerate(headers):
        w = col_widths[col_idx]
        rect = plt.Rectangle((x, y_top - TABLE_HEADER_HEIGHT), w, TABLE_HEADER_HEIGHT,
                            facecolor=TABLE_HEADER_BG, edgecolor=TABLE_BORDER_COLOR, linewidth=0.5)
        ax.add_patch(rect)
        ax.text(x + w/2, y_top - TABLE_HEADER_HEIGHT/2, str(header),
               ha='center', va='center', color=TABLE_HEADER_TEXT,
               fontproperties=font_prop_bold, fontsize=TABLE_HEADER_FONT_SIZE)
        x += w

    # 绘制数据行
    for row_idx, row in enumerate(rows):
        y = y_top - TABLE_HEADER_HEIGHT - (row_idx + 1) * TABLE_ROW_HEIGHT
        bg = TABLE_ROW_BG_EVEN if row_idx % 2 == 0 else TABLE_ROW_BG_ODD
        x = 0
        for col_idx in range(num_cols):
            w = col_widths[col_idx]
            rect = plt.Rectangle((x, y), w, TABLE_ROW_HEIGHT,
                                facecolor=bg, edgecolor=TABLE_BORDER_COLOR, linewidth=0.5)
            ax.add_patch(rect)
            cell_text = str(row[col_idx]) if col_idx < len(row) else ''
            # 替换问题 Unicode 字符
            cell_text = cell_text.replace('\u2717', '\u00d7')
            cell_text = cell_text.replace('\u2718', '\u00d7')
            ax.text(x + w/2, y + TABLE_ROW_HEIGHT/2, cell_text,
                   ha='center', va='center', color=TABLE_TEXT_COLOR,
                   fontproperties=font_prop, fontsize=TABLE_FONT_SIZE)
            x += w

    # 渲染为 PNG，最小 padding
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight", pad_inches=0.02,
                facecolor="white", edgecolor="none")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


# ============================================================
# 图片处理
# ============================================================

def convert_image_to_webp(image_path: str, quality: int = 85) -> bytes:
    """
    将图片转换为 WebP 格式并返回字节数据。

    Args:
        image_path: 原始图片路径
        quality: WebP 压缩质量（1-100）

    Returns:
        WebP 格式的图片字节数据
    """
    if not HAS_PIL:
        with open(image_path, 'rb') as f:
            return f.read()

    img = Image.open(image_path)

    # 确保 RGB 模式
    if img.mode in ('RGBA', 'P'):
        img = img.convert('RGB')

    # 限制最大尺寸（避免 DOCX 过大）
    max_width = 1600
    if img.width > max_width:
        ratio = max_width / img.width
        new_height = int(img.height * ratio)
        img = img.resize((max_width, new_height), Image.LANCZOS)

    buffer = io.BytesIO()
    img.save(buffer, format='WEBP', quality=quality)
    return buffer.getvalue()


# ============================================================
# Markdown 解析
# ============================================================

def parse_markdown_line(line: str) -> Dict:
    """
    解析单行 Markdown，返回解析结果。

    Args:
        line: 单行 Markdown 文本

    Returns:
        解析结果字典，包含 type 和 content
    """
    stripped = line.strip()

    # 空行
    if not stripped:
        return {'type': 'empty', 'content': ''}

    # 分隔线
    if re.match(r'^---+$', stripped):
        return {'type': 'separator', 'content': ''}

    # 标题
    heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
    if heading_match:
        level = len(heading_match.group(1))
        return {'type': f'h{level}', 'content': heading_match.group(2)}

    # 图片引用
    img_match = re.match(r'!\[([^\]]*)\]\(([^\)]+)\)', stripped)
    if img_match:
        return {'type': 'image', 'alt': img_match.group(1), 'src': img_match.group(2)}

    # IMAGE_SLOT 注释
    if stripped.startswith('<!-- IMAGE_SLOT'):
        return {'type': 'image_slot', 'content': stripped}

    # 有序列表
    ol_match = re.match(r'^(\d+)[.、]\s+(.+)$', stripped)
    if ol_match:
        return {'type': 'ordered_list', 'number': int(ol_match.group(1)), 'content': ol_match.group(2)}

    # 无序列表
    ul_match = re.match(r'^[-*+]\s+(.+)$', stripped)
    if ul_match:
        return {'type': 'unordered_list', 'content': ul_match.group(1)}

    # 引用块
    bq_match = re.match(r'^>\s*(.*)$', stripped)
    if bq_match:
        return {'type': 'blockquote', 'content': bq_match.group(1)}

    # 表格行
    if '|' in stripped and stripped.startswith('|'):
        cells = [c.strip() for c in stripped.split('|')[1:-1]]
        if all(re.match(r'^[-:]+$', c) for c in cells):
            return {'type': 'table_separator', 'content': ''}
        return {'type': 'table_row', 'cells': cells}

    # 普通段落
    return {'type': 'paragraph', 'content': stripped}


def clean_markdown_syntax(text: str) -> str:
    """
    清理文本中的 Markdown 语法标记，返回纯文本。

    Args:
        text: 可能含有 Markdown 语法的文本

    Returns:
        清理后的纯文本
    """
    cleaned = text
    for pattern, name, replacement in MARKDOWN_PATTERNS:
        cleaned = re.sub(pattern, replacement, cleaned)
    return cleaned


def detect_markdown_residue(text: str) -> List[Dict]:
    """
    检测文本中的 Markdown 语法残留。

    Args:
        text: 待检测文本

    Returns:
        残留列表
    """
    residues = []
    for pattern, name, _ in MARKDOWN_PATTERNS:
        matches = re.findall(pattern, text)
        if matches:
            residues.append({
                'type': name,
                'count': len(matches),
                'samples': [str(m)[:50] for m in matches[:3]],
            })
    return residues


def detect_placeholders(text: str) -> List[Dict]:
    """
    检测文本中的占位符残留。

    Args:
        text: 待检测文本

    Returns:
        占位符列表
    """
    placeholders = []
    for pattern in PLACEHOLDER_PATTERNS:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            placeholders.append({
                'pattern': pattern,
                'count': len(matches),
                'samples': matches[:3],
            })
    return placeholders


def apply_paragraph_style(paragraph, style_key: str):
    """
    应用段落样式。

    Args:
        paragraph: python-docx Paragraph 对象
        style_key: 样式键名（h1/h2/h3/body/caption）
    """
    if not HAS_DOCX:
        return

    style = STYLE_CONFIG.get(style_key, STYLE_CONFIG['body'])

    for run in paragraph.runs:
        run.font.size = style['size']
        run.font.bold = style['bold']
        run.font.color.rgb = style['color']

    paragraph.paragraph_format.line_spacing = LINE_SPACING
    paragraph.paragraph_format.space_after = Pt(6)


# ============================================================
# DOCX 构建核心
# ============================================================

def _render_table_as_image_to_doc(doc: Document, rows: List[List[str]], temp_dir: str):
    """
    将表格渲染为图片并插入 DOCX 文档。

    替代原来的 _render_table 函数（Word 原生表格），
    改为生成表格图片后以图片形式嵌入，避免媒体平台乱码问题。

    Args:
        doc: Document 对象
        rows: 表格行数据（第一行为表头）
        temp_dir: 临时图片存储目录
    """
    if not rows:
        return

    headers = rows[0]
    data_rows = rows[1:]

    png_data = render_table_as_image(headers, data_rows)

    if png_data is None:
        # Fallback: 以纯文本形式嵌入（当 matplotlib 不可用时）
        p = doc.add_paragraph()
        header_line = " | ".join(headers)
        run = p.add_run(header_line)
        run.font.bold = True
        run.font.size = Pt(10)
        apply_paragraph_style(p, 'body')
        for row in data_rows:
            p = doc.add_paragraph()
            run = p.add_run(" | ".join(row))
            run.font.size = Pt(10)
            apply_paragraph_style(p, 'body')
        return

    # 保存为临时文件并插入
    table_img_path = os.path.join(
        temp_dir, f"table_{id(rows)}_{len(rows)}.png"
    )
    with open(table_img_path, "wb") as f:
        f.write(png_data)

    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(table_img_path, width=Cm(TABLE_INSERT_WIDTH_CM))
    except Exception as e:
        p = doc.add_paragraph()
        run = p.add_run(f"[表格图片嵌入失败: {e}]")
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


def _embed_image(doc: Document, parsed: Dict, images_dir: str, webp_quality: int):
    """
    嵌入图片到 DOCX。

    Args:
        doc: Document 对象
        parsed: 解析后的图片信息
        images_dir: 图片目录
        webp_quality: WebP 质量
    """
    src = parsed.get('src', '')
    alt = parsed.get('alt', '')

    # 查找图片文件
    image_path = os.path.join(images_dir, os.path.basename(src))
    if not os.path.exists(image_path):
        image_path = src

    if os.path.exists(image_path):
        try:
            webp_data = convert_image_to_webp(image_path, webp_quality)
            image_stream = io.BytesIO(webp_data)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(image_stream, width=Inches(min(PAGE_WIDTH_INCHES, 5.5)))
            
            # v6: Image caption is NOT rendered in body to avoid metadiscourse pollution.
            # No caption paragraph is added here.
        except Exception as e:
            p = doc.add_paragraph()
            run = p.add_run(f"[图片嵌入失败: {src} - {e}]")
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    else:
        p = doc.add_paragraph()
        run = p.add_run(f"[图片未找到: {src}]")
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


def clear_article_headers_footers(doc) -> None:
    """Ensure media-ready article DOCX has no internal headers/footers."""
    for section in doc.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                paragraph.text = ""
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell.text = ""


def build_docx(md_lines: List[str], images_dir: str, webp_quality: int = 85) -> Optional[Document]:
    """
    将解析后的 Markdown 行构建为 DOCX 文档。

    ★ 表格以图片形式嵌入，不使用 Word 原生表格。
    最终 DOCX 仅包含纯文本+图片，兼容各类媒体平台。

    Args:
        md_lines: Markdown 文件的行列表
        images_dir: 图片目录路径
        webp_quality: WebP 压缩质量

    Returns:
        python-docx Document 对象
    """
    if not HAS_DOCX:
        print("错误: python-docx 未安装", file=sys.stderr)
        return None

    doc = Document()
    clear_article_headers_footers(doc)

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(11)
    font.name = 'Microsoft YaHei'

    table_buffer = []  # 缓存表格行
    list_counter = 0   # 有序列表计数器
    temp_dir = tempfile.mkdtemp(prefix="frontmind_tables_")

    for line in md_lines:
        parsed = parse_markdown_line(line)

        # 处理表格缓冲
        if parsed['type'] in ('table_row', 'table_separator'):
            if parsed['type'] == 'table_row':
                table_buffer.append(parsed['cells'])
            continue
        elif table_buffer:
            # 表格结束，渲染为图片并插入（非 Word 原生表格）
            _render_table_as_image_to_doc(doc, table_buffer, temp_dir)
            table_buffer = []

        if parsed['type'] == 'empty':
            continue

        elif parsed['type'] == 'separator':
            doc.add_paragraph('').paragraph_format.space_after = Pt(12)

        elif parsed['type'].startswith('h'):
            level = int(parsed['type'][1])
            text = clean_markdown_syntax(parsed['content'])
            style_key = f'h{min(level, 3)}'
            p = doc.add_paragraph()
            run = p.add_run(text)
            apply_paragraph_style(p, style_key)
            list_counter = 0

        elif parsed['type'] == 'image':
            _embed_image(doc, parsed, images_dir, webp_quality)

        elif parsed['type'] == 'image_slot':
            p = doc.add_paragraph()
            run = p.add_run(f"[警告: 未替换的 IMAGE_SLOT: {parsed['content'][:60]}]")
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        elif parsed['type'] == 'ordered_list':
            list_counter += 1
            text = clean_markdown_syntax(parsed['content'])
            p = doc.add_paragraph()
            run = p.add_run(f"{list_counter}. {text}")
            apply_paragraph_style(p, 'body')
            p.paragraph_format.left_indent = Cm(1)

        elif parsed['type'] == 'unordered_list':
            text = clean_markdown_syntax(parsed['content'])
            p = doc.add_paragraph()
            run = p.add_run(f"• {text}")
            apply_paragraph_style(p, 'body')
            p.paragraph_format.left_indent = Cm(1)

        elif parsed['type'] == 'blockquote':
            text = clean_markdown_syntax(parsed['content'])
            p = doc.add_paragraph()
            run = p.add_run(text)
            apply_paragraph_style(p, 'body')
            p.paragraph_format.left_indent = Cm(1.5)
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        elif parsed['type'] == 'paragraph':
            text = clean_markdown_syntax(parsed['content'])
            if text:
                p = doc.add_paragraph()
                run = p.add_run(text)
                apply_paragraph_style(p, 'body')
                list_counter = 0

    # 处理最后的表格缓冲
    if table_buffer:
        _render_table_as_image_to_doc(doc, table_buffer, temp_dir)

    return doc


def strip_article_title_heading(md_lines: List[str]) -> Tuple[List[str], bool]:
    """
    移除正文文件开头的文章标题 H1 和紧随其后的发布副标题。

    v3.7 规则：E4 final.md/final.docx 是不带文章标题的正文文件，
    发布标题只存在于 title_options_reviewed.json，并由交付消息打印。
    """
    lines = list(md_lines)
    idx = 0
    while idx < len(lines) and not lines[idx].strip():
        idx += 1

    stripped = False
    if idx < len(lines) and re.match(r"^#\s+", lines[idx].strip()):
        del lines[idx]
        stripped = True

        # 删除标题后的空行
        while idx < len(lines) and not lines[idx].strip():
            del lines[idx]

        # 删除紧随文章标题的发布副标题/blockquote（章节正文内的引用不受影响）
        if idx < len(lines) and re.match(r"^>\s+", lines[idx].strip()):
            del lines[idx]
            while idx < len(lines) and not lines[idx].strip():
                del lines[idx]

    return lines, stripped


def assemble_with_size_control(md_path: str, images_dir: str,
                                output_path: str, max_mb: float = 10.0) -> Dict:
    """
    组装 DOCX 并控制文件大小。

    如果超过 max_mb，逐步降低 WebP quality。

    Args:
        md_path: Markdown 文件路径
        images_dir: 图片目录路径
        output_path: 输出 DOCX 路径
        max_mb: 最大文件大小（MB）

    Returns:
        组装结果字典
    """
    with open(md_path, 'r', encoding='utf-8') as f:
        md_lines = f.readlines()

    md_lines, article_heading_stripped = strip_article_title_heading(md_lines)

    quality_levels = [85, 70, 55, 40, 25]

    for quality in quality_levels:
        doc = build_docx(md_lines, images_dir, quality)
        if doc is None:
            return {"success": False, "error": "python-docx 未安装"}

        doc.save(output_path)

        size_mb = os.path.getsize(output_path) / (1024 * 1024)

        if size_mb <= max_mb:
            # 检测 Markdown 残留
            full_text = '\n'.join(p.text for p in doc.paragraphs)
            md_residues = detect_markdown_residue(full_text)
            placeholders = detect_placeholders(full_text)

            # 验证无原生表格（纯文本+图片要求）
            native_tables = len(doc.tables)

            return {
                "success": True,
                "webp_quality": quality,
                "file_size_mb": round(size_mb, 2),
                "markdown_residues": md_residues,
                "placeholders": placeholders,
                "native_tables": native_tables,
                "article_heading_stripped": article_heading_stripped,
                "output_path": output_path,
            }

    return {
        "success": False,
        "error": f"文件大小 {size_mb:.2f} MB 仍超过 {max_mb} MB 限制",
        "webp_quality": quality_levels[-1],
        "file_size_mb": round(size_mb, 2),
    }


def main():
    parser = argparse.ArgumentParser(description="FrontMind E4 MD→DOCX 组装器")
    parser.add_argument("--input", required=True, help="输入 Markdown 文件路径")
    parser.add_argument("--images-dir", default="./images/", help="图片目录路径")
    parser.add_argument("--output", required=True, help="输出 DOCX 文件路径")
    parser.add_argument("--max-size", type=float, default=10.0, help="最大文件大小（MB）")

    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"错误: 输入文件不存在: {args.input}", file=sys.stderr)
        sys.exit(1)

    result = assemble_with_size_control(
        args.input, args.images_dir, args.output, args.max_size
    )

    if result["success"]:
        print(f"✅ DOCX 组装成功")
        print(f"  文件: {result['output_path']}")
        print(f"  大小: {result['file_size_mb']} MB")
        print(f"  WebP quality: {result['webp_quality']}")
        print(f"  原生表格数: {result['native_tables']}（应为 0）")
        print(f"  文章标题H1已移除: {result.get('article_heading_stripped', False)}")

        if result["native_tables"] > 0:
            print(f"  ⚠️ 警告: 检测到 {result['native_tables']} 个原生表格，应全部为图片形式")

        if result["markdown_residues"]:
            print(f"  ⚠️ Markdown 残留: {len(result['markdown_residues'])} 处")
            for r in result["markdown_residues"]:
                print(f"    - {r['type']}: {r['count']} 处")

        if result["placeholders"]:
            print(f"  ⚠️ 占位符残留: {len(result['placeholders'])} 处")
            for p in result["placeholders"]:
                print(f"    - {p['pattern']}: {p['count']} 处")
    else:
        print(f"❌ DOCX 组装失败: {result.get('error', '未知错误')}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
